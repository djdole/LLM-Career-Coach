"""Tests for the --generate resume_data workflow: build_source_file_list,
extract_text_from_source_file, build_resume_data_prompt,
validate_resume_data_draft, call_llm_update_resume_data (mocked client),
and the generate_resume_data_draft orchestrator (with the LLM call itself
stubbed)."""

import json
from pathlib import Path
from unittest.mock import MagicMock

import openai
import pytest

import generator


class TestBuildSourceFileList:
    def test_missing_data_dir_returns_empty(self, tmp_path):
        result = generator.build_source_file_list(
            tmp_path / "nope", tmp_path / "kb.json", tmp_path / "kb.draft.json"
        )
        assert result == []

    def test_empty_data_dir_returns_empty(self, tmp_path):
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        assert generator.build_source_file_list(data_dir, tmp_path / "kb.json", tmp_path / "kb.draft.json") == []

    def test_excludes_knowledge_base_file(self, tmp_path):
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        kb_path = data_dir / "resume_data.json"
        kb_path.write_text("{}", encoding="utf-8")
        assert generator.build_source_file_list(data_dir, kb_path, data_dir / "resume_data.json") == []

    def test_excludes_existing_draft_file(self, tmp_path):
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        draft_path = data_dir / "resume_data.json"
        draft_path.write_text("{}", encoding="utf-8")
        assert generator.build_source_file_list(data_dir, data_dir / "resume_data.json", draft_path) == []

    def test_includes_other_files_only(self, tmp_path):
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        kb_path = data_dir / "resume_data.json"
        kb_path.write_text("{}", encoding="utf-8")
        draft_path = data_dir / "resume_data.json"
        new_resume = data_dir / "old_resume.txt"
        new_resume.write_text("some text", encoding="utf-8")
        hidden = data_dir / ".hidden"
        hidden.write_text("nope", encoding="utf-8")
        result = generator.build_source_file_list(data_dir, kb_path, draft_path)
        assert result == [new_resume]

    def test_falls_back_to_unresolved_path_on_resolve_oserror(self, tmp_path, monkeypatch):
        # A permissions issue, filesystem loop, etc. can make Path.resolve()
        # raise OSError; build_source_file_list must fall back to the
        # unresolved path rather than blow up the whole listing.
        data_dir = tmp_path / "data"
        data_dir.mkdir()
        source = data_dir / "old_resume.txt"
        source.write_text("some text", encoding="utf-8")

        original_resolve = Path.resolve

        def flaky_resolve(self, *args, **kwargs):
            raise OSError("simulated resolve failure")

        monkeypatch.setattr(Path, "resolve", flaky_resolve)
        try:
            result = generator.build_source_file_list(data_dir, tmp_path / "kb.json", tmp_path / "draft.json")
        finally:
            monkeypatch.setattr(Path, "resolve", original_resolve)
        assert result == [source]


class TestExtractTextFromSourceFile:
    def test_reads_txt(self, tmp_path):
        p = tmp_path / "notes.txt"
        p.write_text("hello world", encoding="utf-8")
        assert generator.extract_text_from_source_file(p) == "hello world"

    def test_reads_json(self, tmp_path):
        p = tmp_path / "data.json"
        p.write_text('{"a": 1}', encoding="utf-8")
        assert generator.extract_text_from_source_file(p) == '{"a": 1}'

    def test_reads_xml(self, tmp_path):
        p = tmp_path / "data.xml"
        p.write_text("<root><a>1</a></root>", encoding="utf-8")
        assert "root" in generator.extract_text_from_source_file(p)

    def test_unsupported_extension_returns_empty(self, tmp_path):
        p = tmp_path / "data.exe"
        p.write_bytes(b"\x00\x01")
        assert generator.extract_text_from_source_file(p) == ""

    def test_reads_docx_paragraphs_and_tables(self, tmp_path):
        from docx import Document

        p = tmp_path / "resume.docx"
        doc = Document()
        doc.add_paragraph("Jane Doe, Software Engineer")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Skill"
        table.rows[0].cells[1].text = "Python"
        doc.save(str(p))

        text = generator.extract_text_from_source_file(p)
        assert "Jane Doe, Software Engineer" in text
        assert "Skill | Python" in text

    def test_reads_pdf_text(self, tmp_path):
        from reportlab.pdfgen import canvas

        p = tmp_path / "resume.pdf"
        c = canvas.Canvas(str(p))
        c.drawString(72, 720, "Jane Doe, Software Engineer")
        c.save()

        text = generator.extract_text_from_source_file(p)
        assert "Jane Doe, Software Engineer" in text

    def test_unreadable_file_returns_empty_and_logs(self, tmp_path, capsys):
        # A .docx-suffixed file that isn't actually a valid docx makes
        # python-docx raise; extract_text_from_source_file must catch
        # that (any Exception), log it, and return "" rather than
        # propagate and abort the whole resume_data run over one bad file.
        p = tmp_path / "corrupt.docx"
        p.write_bytes(b"not a real docx file")
        assert generator.extract_text_from_source_file(p) == ""
        assert "Could not read corrupt.docx" in capsys.readouterr().err


class TestValidateResumeDataDraft:
    def test_accepts_valid_new_draft(self):
        draft = {"personal_info": {}, "education": [], "skills": {}, "work_experience": []}
        generator.validate_resume_data_draft(draft, None)  # no exception

    def test_rejects_missing_required_section(self):
        draft = {"personal_info": {}, "education": [], "skills": {}}
        with pytest.raises(ValueError, match="missing required section"):
            generator.validate_resume_data_draft(draft, None)

    def test_rejects_non_dict(self):
        with pytest.raises(ValueError):
            generator.validate_resume_data_draft(["not", "a", "dict"], None)

    def test_rejects_dropped_top_level_section_on_update(self, sample_kb):
        draft = {"personal_info": {}, "education": [], "skills": {}, "work_experience": []}
        with pytest.raises(ValueError, match="non-destructive"):
            generator.validate_resume_data_draft(draft, sample_kb)

    def test_accepts_update_that_keeps_all_sections(self, sample_kb):
        draft = dict(sample_kb)
        draft["skills"] = {**draft["skills"], "new_category": ["Something"]}
        generator.validate_resume_data_draft(draft, sample_kb)  # no exception


class TestBuildResumeDataPrompt:
    def test_from_scratch_build_has_no_existing_kb_section(self):
        prompt = generator.build_resume_data_prompt(None, {"resume.txt": "Jane Doe, SWE"})
        assert "BRAND NEW resume knowledge base" in prompt
        assert "=== EXISTING KNOWLEDGE BASE" not in prompt
        assert "Jane Doe, SWE" in prompt

    def test_update_includes_existing_kb_and_non_destructive_instructions(self, sample_kb):
        prompt = generator.build_resume_data_prompt(sample_kb, {"new_job.txt": "Started at Beta Corp"})
        assert "NON-DESTRUCTIVE" in prompt
        assert "=== EXISTING KNOWLEDGE BASE" in prompt
        assert "Started at Beta Corp" in prompt
        # The existing KB itself is embedded as JSON for the model to see.
        assert sample_kb["personal_info"]["full_name"] in prompt

    def test_includes_every_source_file_under_its_own_heading(self):
        prompt = generator.build_resume_data_prompt(None, {"a.txt": "Content A", "b.txt": "Content B"})
        assert "=== SOURCE FILE: a.txt ===\nContent A" in prompt
        assert "=== SOURCE FILE: b.txt ===\nContent B" in prompt


class TestCallLlmUpdateResumeData:
    def _make_response(self, content):
        response = MagicMock()
        response.choices = [MagicMock()]
        response.choices[0].message.content = content
        return response

    def test_builds_new_kb_on_first_attempt(self):
        new_kb = {"personal_info": {}, "education": [], "skills": {}, "work_experience": []}
        client = MagicMock()
        client.chat.completions.create.return_value = self._make_response(json.dumps(new_kb))
        result = generator.call_llm_update_resume_data(client, None, {"resume.txt": "Jane Doe"})
        assert result == new_kb
        assert client.chat.completions.create.call_count == 1

    def test_updates_existing_kb_on_first_attempt(self, sample_kb):
        client = MagicMock()
        client.chat.completions.create.return_value = self._make_response(json.dumps(sample_kb))
        result = generator.call_llm_update_resume_data(client, sample_kb, {"new_job.txt": "New role"})
        assert result == sample_kb

    def test_retries_on_unparsable_json_then_succeeds(self):
        new_kb = {"personal_info": {}, "education": [], "skills": {}, "work_experience": []}
        client = MagicMock()
        client.chat.completions.create.side_effect = [
            self._make_response("not json at all"),
            self._make_response(json.dumps(new_kb)),
        ]
        result = generator.call_llm_update_resume_data(client, None, {"resume.txt": "Jane Doe"})
        assert result == new_kb
        assert client.chat.completions.create.call_count == 2

    def test_retries_when_draft_drops_existing_section(self, sample_kb):
        # First response is valid JSON but fails validate_resume_data_draft
        # (drops a top-level section from the existing KB) -> retry.
        bad_draft = {"personal_info": {}, "education": [], "skills": {}, "work_experience": []}
        good_draft = dict(sample_kb)
        client = MagicMock()
        client.chat.completions.create.side_effect = [
            self._make_response(json.dumps(bad_draft)),
            self._make_response(json.dumps(good_draft)),
        ]
        result = generator.call_llm_update_resume_data(client, sample_kb, {"new_job.txt": "New role"})
        assert result == good_draft
        assert client.chat.completions.create.call_count == 2

    def test_exits_after_exhausting_retries(self):
        client = MagicMock()
        client.chat.completions.create.return_value = self._make_response("still not valid json")
        with pytest.raises(SystemExit) as exc_info:
            generator.call_llm_update_resume_data(client, None, {"resume.txt": "Jane Doe"})
        assert exc_info.value.code == 1
        assert client.chat.completions.create.call_count == 2

    def test_exits_on_connection_error(self):
        client = MagicMock()
        client.chat.completions.create.side_effect = openai.APIConnectionError(
            request=MagicMock(), message="unreachable"
        )
        with pytest.raises(SystemExit) as exc_info:
            generator.call_llm_update_resume_data(client, None, {"resume.txt": "Jane Doe"})
        assert exc_info.value.code == 1

    def test_exits_on_api_status_error(self):
        client = MagicMock()
        client.chat.completions.create.side_effect = openai.APIStatusError(
            message="server error",
            response=MagicMock(status_code=500),
            body={"error": "server error"},
        )
        with pytest.raises(SystemExit) as exc_info:
            generator.call_llm_update_resume_data(client, None, {"resume.txt": "Jane Doe"})
        assert exc_info.value.code == 1


@pytest.fixture
def resume_data_env(tmp_path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("LITELLM_BASE_URL", "http://example.com")
    monkeypatch.setenv("LITELLM_API_KEY", "secret")
    return tmp_path


class TestGenerateResumeDataDraft:
    def test_noop_when_data_env_unset(self, resume_data_env, monkeypatch):
        monkeypatch.delenv("DATA", raising=False)
        settings = generator.load_file_location_settings()
        generator.generate_resume_data_draft(client=None, s=settings)
        assert not (resume_data_env / "data").exists()

    def test_noop_when_data_folder_missing(self, resume_data_env, monkeypatch):
        monkeypatch.setenv("DATA", "data")
        settings = generator.load_file_location_settings()
        generator.generate_resume_data_draft(client=None, s=settings)
        assert not (resume_data_env / "data" / "resume_data.json").exists()

    def test_noop_when_data_folder_empty(self, resume_data_env, monkeypatch):
        monkeypatch.setenv("DATA", "data")
        (resume_data_env / "data").mkdir()
        settings = generator.load_file_location_settings()
        generator.generate_resume_data_draft(client=None, s=settings)
        assert not (resume_data_env / "data" / "resume_data.json").exists()

    def test_noop_when_only_knowledge_base_present(self, resume_data_env, monkeypatch, sample_kb):
        monkeypatch.setenv("DATA", "data")
        data_dir = resume_data_env / "data"
        data_dir.mkdir()
        (data_dir / "resume_data.json").write_text(json.dumps(sample_kb), encoding="utf-8")
        settings = generator.load_file_location_settings()
        generator.generate_resume_data_draft(client=None, s=settings)
        assert (data_dir / "resume_data.json").exists()

    def test_builds_new_draft_when_knowledge_base_missing(self, resume_data_env, monkeypatch):
        monkeypatch.setenv("DATA", "data")
        data_dir = resume_data_env / "data"
        data_dir.mkdir()
        source = data_dir / "old_resume.txt"
        source.write_text("Jane Doe, Software Engineer at Acme", encoding="utf-8")
        settings = generator.load_file_location_settings()

        new_kb = {"personal_info": {"full_name": "Jane Doe"}, "education": [], "skills": {}, "work_experience": []}
        monkeypatch.setattr(
            generator, "call_llm_update_resume_data",
            lambda client, existing_kb, source_texts: new_kb if existing_kb is None else (_ for _ in ()).throw(AssertionError("expected no existing kb")),
        )

        generator.generate_resume_data_draft(client=object(), s=settings)

        draft_path = data_dir / "resume_data.json"
        assert json.loads(draft_path.read_text(encoding="utf-8")) == new_kb
        assert not source.exists()  # consumed source file removed

    def test_updates_draft_non_destructively_when_knowledge_base_exists(self, resume_data_env, monkeypatch, sample_kb):
        monkeypatch.setenv("DATA", "data")
        data_dir = resume_data_env / "data"
        data_dir.mkdir()
        (data_dir / "resume_data.json").write_text(json.dumps(sample_kb), encoding="utf-8")
        source = data_dir / "new_job.txt"
        source.write_text("Started a new role at Beta Corp", encoding="utf-8")
        settings = generator.load_file_location_settings()

        updated_kb = dict(sample_kb)
        captured = {}

        def fake_call(client, existing_kb, source_texts):
            captured["existing_kb"] = existing_kb
            captured["source_texts"] = source_texts
            return updated_kb

        monkeypatch.setattr(generator, "call_llm_update_resume_data", fake_call)

        generator.generate_resume_data_draft(client=object(), s=settings)

        assert captured["existing_kb"] == sample_kb
        assert "new_job.txt" in captured["source_texts"]
        draft_path = data_dir / "resume_data.json"
        assert json.loads(draft_path.read_text(encoding="utf-8")) == updated_kb
        assert not source.exists()  # consumed
        assert (data_dir / "resume_data.json").exists()  # knowledge base itself untouched

    def test_leaves_existing_draft_alone_when_no_new_sources(self, resume_data_env, monkeypatch, sample_kb):
        monkeypatch.setenv("DATA", "data")
        data_dir = resume_data_env / "data"
        data_dir.mkdir()
        (data_dir / "resume_data.json").write_text(json.dumps(sample_kb), encoding="utf-8")
        stale_draft = data_dir / "resume_data.json"
        stale_draft.write_text('{"stale": true}', encoding="utf-8")
        settings = generator.load_file_location_settings()

        generator.generate_resume_data_draft(client=None, s=settings)

        # No new source files besides the stale draft/kb -> untouched.
        assert json.loads(stale_draft.read_text(encoding="utf-8")) == {"stale": True}

    def test_noop_when_source_files_have_no_extractable_text(self, resume_data_env, monkeypatch):
        monkeypatch.setenv("DATA", "data")
        data_dir = resume_data_env / "data"
        data_dir.mkdir()
        # An unsupported extension -> extract_text_from_source_file
        # returns "" -> source_texts ends up empty -> skip, don't call
        # the LLM at all.
        (data_dir / "notes.exe").write_bytes(b"\x00\x01")
        settings = generator.load_file_location_settings()

        called = []
        monkeypatch.setattr(
            generator, "call_llm_update_resume_data", lambda client, existing_kb, source_texts: called.append(True)
        )
        generator.generate_resume_data_draft(client=None, s=settings)

        assert called == []
        assert not (data_dir / "resume_data.json").exists()

    def test_logs_but_does_not_fail_when_a_consumed_source_file_cant_be_removed(
        self, resume_data_env, monkeypatch, capsys
    ):
        monkeypatch.setenv("DATA", "data")
        data_dir = resume_data_env / "data"
        data_dir.mkdir()
        source = data_dir / "old_resume.txt"
        source.write_text("Jane Doe, Software Engineer at Acme", encoding="utf-8")
        settings = generator.load_file_location_settings()

        new_kb = {"personal_info": {"full_name": "Jane Doe"}, "education": [], "skills": {}, "work_experience": []}
        monkeypatch.setattr(generator, "call_llm_update_resume_data", lambda client, existing_kb, source_texts: new_kb)

        def flaky_unlink(self, *args, **kwargs):
            raise OSError("simulated permission error")

        monkeypatch.setattr(Path, "unlink", flaky_unlink)

        generator.generate_resume_data_draft(client=object(), s=settings)

        # The draft still gets written even though cleanup partially failed.
        assert json.loads((data_dir / "resume_data.json").read_text(encoding="utf-8")) == new_kb
        assert source.exists()  # removal failed, so it's still there
        assert "Could not remove consumed source file" in capsys.readouterr().err