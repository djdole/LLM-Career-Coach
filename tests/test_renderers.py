"""Tests for the render_* functions: plain text, Markdown, DOCX, and PDF
output for both the resume and the cover letter."""

from docx import Document

import generator


# --- Plain text / Markdown --------------------------------------------

class TestRenderResumeTxt:
    def test_contains_name_and_contact(self, sample_resume_dict):
        out = generator.render_resume_txt(sample_resume_dict)
        assert out.startswith("Jane Q. Doe\n")
        assert sample_resume_dict["contact_line"] in out

    def test_strips_em_dashes(self, sample_resume_dict):
        out = generator.render_resume_txt(sample_resume_dict)
        assert "\u2014" not in out
        assert "does things" in out

    def test_includes_bullets_with_bullet_char(self, sample_resume_dict):
        out = generator.render_resume_txt(sample_resume_dict)
        assert f"{generator.BULLET_CHAR} Did a thing." in out

    def test_team_context_appended_to_company_line(self, sample_resume_dict):
        out = generator.render_resume_txt(sample_resume_dict)
        assert "Acme Corp \u00b7 Team Widgets" in out

    def test_job_without_team_context_has_bare_company_line(self, sample_resume_dict):
        out = generator.render_resume_txt(sample_resume_dict)
        lines = out.split("\n")
        assert "Beta Inc" in lines

    def test_includes_education(self, sample_resume_dict):
        out = generator.render_resume_txt(sample_resume_dict)
        assert "State University (2015)" in out


class TestRenderResumeMd:
    def test_uses_markdown_headings(self, sample_resume_dict):
        out = generator.render_resume_md(sample_resume_dict)
        assert out.startswith("# Jane Q. Doe\n")
        assert "## Summary" in out
        assert "## Work Experience" in out
        assert "## Education" in out

    def test_skills_heading_is_title_cased(self, sample_resume_dict):
        out = generator.render_resume_md(sample_resume_dict)
        assert "## Core Technical Skills" in out

    def test_strips_em_dashes(self, sample_resume_dict):
        out = generator.render_resume_md(sample_resume_dict)
        assert "\u2014" not in out

    def test_job_header_uses_pipe_between_title_and_dates(self, sample_resume_dict):
        out = generator.render_resume_md(sample_resume_dict)
        assert "### Software Engineer | 2020-01 - 2023-01" in out


class TestRenderCoverLetterTxtAndMd:
    def test_txt_strips_em_dashes(self, sample_cover_letter_dict):
        out = generator.render_cover_letter_txt(sample_cover_letter_dict)
        assert "\u2014" not in out
        assert "truly" in out

    def test_md_matches_txt(self, sample_cover_letter_dict):
        assert generator.render_cover_letter_md(sample_cover_letter_dict) == \
            generator.render_cover_letter_txt(sample_cover_letter_dict)


# --- DOCX ----------------------------------------------------------------

class TestRenderResumeDocx:
    def test_writes_a_nonempty_file(self, tmp_path, sample_resume_dict):
        out_path = tmp_path / "resume.docx"
        generator.render_resume_docx(sample_resume_dict, out_path)
        assert out_path.exists()
        assert out_path.stat().st_size > 0

    def test_document_contains_expected_text(self, tmp_path, sample_resume_dict):
        out_path = tmp_path / "resume.docx"
        generator.render_resume_docx(sample_resume_dict, out_path)
        doc = Document(str(out_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Jane Q. Doe" in all_text
        assert "does things" in all_text  # em dash stripped, comma-joined
        assert "\u2014" not in all_text

    def test_respects_body_pt_font_size(self, tmp_path, sample_resume_dict):
        out_path = tmp_path / "resume.docx"
        generator.render_resume_docx(sample_resume_dict, out_path, body_pt=9)
        doc = Document(str(out_path))
        assert doc.styles["Normal"].font.size.pt == 9


class TestRenderCoverLetterDocx:
    def test_writes_a_nonempty_file(self, tmp_path, sample_cover_letter_dict):
        out_path = tmp_path / "cl.docx"
        generator.render_cover_letter_docx(sample_cover_letter_dict, out_path)
        assert out_path.exists()
        assert out_path.stat().st_size > 0

    def test_paragraphs_split_on_blank_lines(self, tmp_path, sample_cover_letter_dict):
        out_path = tmp_path / "cl.docx"
        generator.render_cover_letter_docx(sample_cover_letter_dict, out_path)
        doc = Document(str(out_path))
        expected_paragraph_count = len(sample_cover_letter_dict["body"].split("\n\n"))
        assert len(doc.paragraphs) == expected_paragraph_count


# --- PDF -------------------------------------------------------------------

class TestRenderResumePdf:
    def test_writes_a_nonempty_file_and_fits_default_tier(self, tmp_path, sample_resume_dict):
        out_path = tmp_path / "resume.pdf"
        pages, body_pt = generator.render_resume_pdf(sample_resume_dict, out_path)
        assert out_path.exists()
        assert out_path.stat().st_size > 0
        assert pages == 1
        assert body_pt == generator.PDF_TIERS[0][0]

    def test_falls_back_to_smallest_tier_when_content_is_huge(self, tmp_path, sample_resume_dict):
        big = dict(sample_resume_dict)
        big["work_experience"] = [
            {
                "title": f"Engineer {i}",
                "company": f"Company {i} Incorporated",
                "date_range": "2020-01 - 2023-01",
                "team_context": "Some team context here",
                "bullets": [f"Did thing number {j} in a fairly long descriptive sentence." for j in range(8)],
            }
            for i in range(10)
        ]
        out_path = tmp_path / "big.pdf"
        pages, body_pt = generator.render_resume_pdf(big, out_path, max_pages=2)
        assert out_path.exists()
        # Content this large can't fit in 2 pages even at the smallest tier,
        # so it should fall through to the last tier and still return cleanly.
        assert body_pt == generator.PDF_TIERS[-1][0]
        assert pages > 2

    def test_job_without_team_context_still_renders(self, tmp_path, sample_resume_dict):
        out_path = tmp_path / "resume.pdf"
        # sample_resume_dict's second job has an empty team_context already;
        # this just confirms rendering doesn't crash on that branch.
        pages, _ = generator.render_resume_pdf(sample_resume_dict, out_path)
        assert pages >= 1


class TestRenderCoverLetterPdf:
    def test_writes_a_nonempty_file(self, tmp_path, sample_cover_letter_dict):
        out_path = tmp_path / "cl.pdf"
        generator.render_cover_letter_pdf(sample_cover_letter_dict, out_path)
        assert out_path.exists()
        assert out_path.stat().st_size > 0
