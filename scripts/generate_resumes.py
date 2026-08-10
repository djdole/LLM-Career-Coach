#!/usr/bin/env python3
"""
Generates baseline (no-specific-JD) SDE and SDET resumes, plus a cover
letter, from data/resume_data.json, in pdf/docx/txt/md/json formats.

Uses a self-hosted LiteLLM proxy (in front of Ollama, per that stack's
docker-compose.yml) rather than a paid hosted API, so this never spends
API credits and never fails due to account balance.

Tailoring a resume to a *specific* job posting is a different task (it
requires selecting/adapting content to a JD) and stays a separate,
manual, chat-based workflow -- see data/resume_data.json's own
generation_workflow_for_llm for that path. This script does not do that.

Usage:
    python scripts/generate_resumes.py --data data/resume_data.json --out generated/
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

import openai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

VARIANTS = ["SDE", "SDET"]

# Points at a self-hosted LiteLLM proxy instead of the Anthropic API. This
# stack runs LiteLLM in front of Ollama specifically as an OpenAI-compatible
# gateway (see docker-compose.yml), so we call LiteLLM directly rather than
# Open WebUI's own API layer. Avoids paid-API token usage entirely;
# generation runs against whatever local model Ollama has loaded.
#
# Required repo configuration (Settings -> Secrets and variables -> Actions):
#   Secrets:   LITELLM_BASE_URL   e.g. https://litellm.example.com
#                                 (points at the LiteLLM container's port,
#                                 ${LITELLM_PORT} in docker-compose.yml --
#                                 NOT Open WebUI's UI port)
#              LITELLM_API_KEY    same value as LITELLM_MASTER_KEY in that
#                                 stack's .env
#   Variables: LITELLM_MODEL      the model string LiteLLM proxies to, e.g.
#                                 "ollama/llama3.1:70b" -- NOT sensitive,
#                                 so it's a repo Variable rather than a
#                                 Secret.
#   Optional:  OLLAMA_NUM_CTX     context window size override (default
#                                 16384) -- lower if your GPU can't hold
#                                 that much context for the model in use.
MODEL = os.environ.get("LITELLM_MODEL", "ollama/llama3.1:70b")

# Maps the knowledge base's snake_case skill category keys to display
# labels matching the existing hand-written resumes' style, fed to the
# model as already-formatted so it only has to copy them, not invent
# formatting for them.
CATEGORY_LABELS = {
    "languages": "Languages",
    "apis_and_web_servers": "APIs & Web Servers",
    "test_automation_frameworks": "Test Automation Frameworks",
    "frontend_and_mobile_testing": "Frontend & Mobile Testing",
    "unit_integration_testing": "Unit & Integration Testing",
    "test_management_and_planning": "Test Management & Planning",
    "performance_testing": "Performance Testing",
    "ci_cd_and_devops": "CI/CD & DevOps",
    "databases_and_query_languages": "Databases & Query Languages",
    "version_control": "Version Control",
    "debugging_and_diagnostics": "Debugging & Diagnostics",
    "virtualization_and_infra": "Virtualization & Infra",
    "developer_tools_and_ides": "Developer Tools & IDEs",
    "ai_and_automation_tooling": "AI & Automation Tooling",
    "methodologies": "Methodologies",
    "collaboration_and_docs": "Collaboration & Docs",
    "monitoring_and_incident_mgmt": "Monitoring / Incident Mgmt",
}

SKILLS_HEADING_BY_VARIANT = {"SDE": "CORE TECHNICAL SKILLS", "SDET": "CORE SDET SKILLS"}
BULLET_CHAR = "\u25cf"  # "●", matches the existing hand-written resumes' style
EM_DASH = "\u2014"


def compute_job_column_widths(work_experience: list, body_pt: float, total_pt: float, min_title_pt: float = 130) -> tuple:
    """
    Sizes the title/employer/date columns from the ACTUAL text in this
    resume at this font size, instead of fixed percentages -- fixed splits
    silently overflow whenever a company name (e.g. 'Cosworth Tech Inc. /
    MAHLE Powertrain LLC') is longer than whatever guess produced the
    split, forcing it to wrap mid-name. Measured with reportlab's
    Helvetica-Bold metrics as a stand-in font for both PDF and DOCX --
    DOCX's actual font differs slightly, but widths are close enough that
    this still prevents wrapping in practice.

    Employer and date each get exactly what their longest actual value
    needs (plus a small buffer); title gets whatever's left, floored at
    min_title_pt so a long employer name can't crush the title column to
    nothing.
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth
    pad = 8
    employer_pt = max(stringWidth(j["company"], "Helvetica-Bold", body_pt) for j in work_experience) + pad
    date_pt = max(stringWidth(j["date_range"], "Helvetica-Bold", body_pt) for j in work_experience) + pad
    title_pt = max(total_pt - employer_pt - date_pt, min_title_pt)
    return title_pt, employer_pt, date_pt


def strip_em_dashes(text: str) -> str:
    """Belt-and-suspenders fallback behind the never_use_em_dash rule the
    model is given -- not a substitute for the model actually following it."""
    return text.replace(EM_DASH, ",")


# --- Prompt construction -------------------------------------------------

def build_baseline_context(kb: dict, variant: str) -> dict:
    """
    Trims the full knowledge base down to only what the BASELINE (no-JD)
    path needs for ONE variant (SDE or SDET), per generation_workflow_for_llm
    step 0's fallback.

    Trimming matters for two reasons: (1) token budget -- the full KB is
    ~10k tokens alone, likely exceeding a local model's context window
    unless num_ctx is raised (see call_llm); (2) signal-to-noise -- the
    full KB includes JD-tailoring-only fields (bullet variants, per-bullet
    themes/skills tags, the other variant's title, cover letter JD-specific
    building blocks) that have shown up verbatim in bad output before this
    trimming existed -- fewer irrelevant JSON shapes nearby means less for
    a smaller model to latch onto instead of the requested schema.

    Skill category labels are pre-formatted here (see CATEGORY_LABELS) so
    the model only has to copy them, not invent formatting.
    """
    rules = kb["meta"]["output_rules"]
    summary = kb["summary_variants"].get(variant) or kb["summary_variants"]["SDE"]
    skills = [
        {"category": CATEGORY_LABELS.get(k, k.replace("_", " ").title()), "items": v}
        for k, v in kb["skills"].items() if isinstance(v, list)
    ]
    work_experience = [
        {
            "title": job["title_by_variant"].get(variant) or next(iter(job["title_by_variant"].values())),
            "company": job["company"],
            "team_context": job.get("team_context", ""),
            "date_range": f"{job['start_date']} - {job['end_date']}",
            "bullets": [b["text"] for b in job["bullets"] if not b["id"].endswith(("_alt", "_variant"))],
        }
        for job in kb["work_experience"]
    ]
    # The raw KB uses the key "graduation date" (with a space); the output
    # schema calls it "date". Renaming it here -- rather than relying on the
    # model to perform that rename -- is what actually fixed education[]
    # entries coming back without a date at all.
    education = [
        {"degree": ed["degree"], "institution": ed["institution"], "date": ed["graduation date"]}
        for ed in kb["education"]
    ]
    return {
        "output_rules": {
            "never_fabricate": rules["never_fabricate"],
            "never_use_em_dash": rules["never_use_em_dash"],
        },
        "personal_info": kb["personal_info"],
        "education": education,
        "summary": summary,
        "skills_heading": SKILLS_HEADING_BY_VARIANT[variant],
        "skills": skills,
        "work_experience": work_experience,
        "cover_letter_generic_template": kb["cover_letter_building_blocks"]["generic_fallback_template"],
    }


def build_system_prompt(kb: dict, variant: str) -> str:
    """
    Schema is placed AFTER the source data and includes a filled-in
    example (not just types), since smaller/local models otherwise tend to
    echo the nearest JSON shape they've seen instead of the requested one.
    """
    context = build_baseline_context(kb, variant)
    return (
        f"You are generating a BASELINE {variant} resume (no specific job "
        "description provided) and a cover letter, for the candidate "
        "described below. Use the summary, skills, and skills_heading "
        "exactly as given, all work_experience bullets in the order given, "
        "and cover_letter_generic_template for the cover letter (lightly "
        "adapt it, keep 'Dear Hiring Manager'). Follow output_rules "
        "exactly, especially never_fabricate and never_use_em_dash.\n\n"
        "=== CANDIDATE DATA (read this for content only; its field names "
        "like team_context, date_range, and cover_letter_generic_template "
        "belong to THIS input and must NOT appear in your answer) ===\n"
        + json.dumps(context, indent=2)
        + "\n\n=== YOUR TASK ===\n"
        "Using only the content above, produce ONE JSON object with EXACTLY "
        "these top-level keys: \"resume\" and \"cover_letter\". Nothing else. "
        "No prose before or after it, no markdown code fences, no headings, "
        "no explanation of what you did.\n\n"
        "Here is a SHAPE EXAMPLE with placeholder values, showing the exact "
        "keys your answer must use (do not copy these placeholder values, "
        "replace them with real content from the candidate data above):\n"
        "{\n"
        '  "resume": {\n'
        '    "name": "Full Name",\n'
        '    "contact_line": "email \u00b7 phone \u00b7 linkedin \u00b7 portfolio",\n'
        '    "skills_heading": "CORE TECHNICAL SKILLS",\n'
        '    "summary": "2-4 sentence professional summary",\n'
        '    "skills": [{"category": "Languages", "items": ["Python", "Go"]}],\n'
        '    "work_experience": [\n'
        "      {\n"
        '        "title": "Job Title", "company": "Company Name",\n'
        '        "date_range": "2020 - 2023", "team_context": "Team X",\n'
        '        "bullets": ["Accomplishment one.", "Accomplishment two."]\n'
        "      }\n"
        "    ],\n"
        '    "education": [{"degree": "B.S. Computer Science", "institution": "Some University", "date": "2010"}]\n'
        "  },\n"
        '  "cover_letter": {"body": "Full cover letter text, paragraphs separated by a blank line."}\n'
        "}\n\n"
        "Respond with ONLY that JSON object. Your entire response must start "
        "with { and end with }."
    )


REQUIRED_RESUME_KEYS = {"name", "contact_line", "summary", "skills", "work_experience", "education"}
REQUIRED_JOB_KEYS = {"title", "company", "date_range", "bullets"}
REQUIRED_EDUCATION_KEYS = {"degree", "institution", "date"}


def extract_json_object(raw: str) -> str:
    """Pulls the first balanced {...} block out of a string that may
    contain markdown fences and/or prose before/after it."""
    start = raw.find("{")
    if start == -1:
        raise ValueError("No '{' found in model output.")
    depth = 0
    for i, ch in enumerate(raw[start:], start=start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return raw[start:i + 1]
    raise ValueError("No balanced closing '}' found in model output.")


def validate_content(content: dict) -> None:
    """Checks both top-level keys AND the nested shape of each
    work_experience/education entry -- the renderers assume every entry
    has these fields, so a model response that's missing one needs to be
    caught HERE (and retried) rather than crashing deep inside rendering."""
    if not isinstance(content, dict) or "resume" not in content or "cover_letter" not in content:
        raise ValueError("Missing top-level 'resume' and/or 'cover_letter' keys.")
    resume = content["resume"]
    missing = REQUIRED_RESUME_KEYS - set(resume.keys())
    if missing:
        raise ValueError(f"resume is missing required keys: {missing}")
    for i, job in enumerate(resume["work_experience"]):
        job_missing = REQUIRED_JOB_KEYS - set(job.keys())
        if job_missing:
            raise ValueError(f"work_experience[{i}] is missing required keys: {job_missing}")
    for i, ed in enumerate(resume["education"]):
        ed_missing = REQUIRED_EDUCATION_KEYS - set(ed.keys())
        if ed_missing:
            raise ValueError(f"education[{i}] is missing required keys: {ed_missing}")
    if "body" not in content["cover_letter"]:
        raise ValueError("cover_letter is missing required key: 'body'")


def call_llm(kb: dict, variant: str) -> dict:
    base_url = os.environ.get("LITELLM_BASE_URL")
    api_key = os.environ.get("LITELLM_API_KEY")
    if not base_url or not api_key:
        print(
            "LITELLM_BASE_URL and/or LITELLM_API_KEY are not set. "
            "Set them as repo secrets (Settings -> Secrets and variables -> "
            "Actions) before running this workflow. LITELLM_API_KEY should "
            "match LITELLM_MASTER_KEY in that stack's .env.",
            file=sys.stderr,
        )
        sys.exit(1)

    client = openai.OpenAI(base_url=base_url.rstrip("/") + "/v1", api_key=api_key)
    system_prompt = build_system_prompt(kb, variant)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Generate the baseline {variant} resume and cover letter JSON now."},
    ]

    est_input_tokens = len(system_prompt) // 4
    print(f"[{variant}] Prompt size estimate: ~{est_input_tokens} input tokens.", file=sys.stderr)

    max_output_tokens = int(os.environ.get("LITELLM_MAX_TOKENS", "10000"))
    num_ctx = int(os.environ.get("OLLAMA_NUM_CTX", "16384"))

    last_error = None
    for attempt in range(2):  # one retry with a corrective follow-up if the first reply is malformed
        try:
            try:
                response = client.chat.completions.create(
                    model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=180,
                    response_format={"type": "json_object"},
                    extra_body={"options": {"num_ctx": num_ctx}},
                    messages=messages,
                )
            except openai.BadRequestError:
                response = client.chat.completions.create(
                    model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=180,
                    extra_body={"options": {"num_ctx": num_ctx}},
                    messages=messages,
                )
        except openai.APIConnectionError as e:
            print(
                f"Could not reach LiteLLM at {base_url}. Is the stack running and "
                "reachable from this runner? (Hosted GitHub runners cannot reach "
                "a private home LAN instance unless it's tunneled/exposed, or "
                "unless this workflow runs on a self-hosted runner on the same "
                f"network.) Underlying error: {e}",
                file=sys.stderr,
            )
            sys.exit(1)
        except openai.APIStatusError as e:
            print(
                f"LiteLLM returned an error (HTTP {e.status_code}): {e.message}. "
                "Check LITELLM_API_KEY matches LITELLM_MASTER_KEY and that "
                "LITELLM_MODEL matches the --model argument LiteLLM was started with.",
                file=sys.stderr,
            )
            sys.exit(1)

        raw = response.choices[0].message.content or ""
        try:
            content = json.loads(extract_json_object(raw))
            validate_content(content)
            return content
        except (ValueError, json.JSONDecodeError) as e:
            last_error = e
            print(f"[{variant}] Attempt {attempt + 1}: malformed output ({e}). Raw:\n{raw}\n", file=sys.stderr)
            messages.append({"role": "assistant", "content": raw})
            messages.append({
                "role": "user",
                "content": (
                    f"That response was invalid: {e}. Respond again with ONLY the "
                    "JSON object described earlier -- no prose, no markdown fences, "
                    "no commentary, and use exactly the keys resume/cover_letter."
                ),
            })

    print(f"[{variant}] Model failed to produce valid JSON after 2 attempts. Last error: {last_error}", file=sys.stderr)
    sys.exit(1)


# --- Plain text / Markdown -------------------------------------------------

def render_resume_txt(r: dict) -> str:
    lines = [r["name"], r["contact_line"], "", "SUMMARY", r["summary"], "", r["skills_heading"]]
    for group in r["skills"]:
        lines.append(f"{group['category']}: {', '.join(group['items'])}")
    lines += ["", "WORK EXPERIENCE"]
    for job in r["work_experience"]:
        lines.append(f"{job['title']} {job['date_range']}")
        line2 = job["company"] + (f" \u00b7 {job['team_context']}" if job.get("team_context") else "")
        lines.append(line2)
        for b in job["bullets"]:
            lines.append(f"{BULLET_CHAR} {b}")
        lines.append("")
    lines.append("EDUCATION")
    for ed in r["education"]:
        lines.append(f"{ed['degree']}")
        lines.append(f"{ed['institution']} ({ed['date']})")
    return strip_em_dashes("\n".join(lines))


def render_resume_md(r: dict) -> str:
    lines = [f"# {r['name']}", r["contact_line"], "", "## Summary", r["summary"], "", f"## {r['skills_heading'].title()}"]
    for group in r["skills"]:
        lines.append(f"- **{group['category']}:** {', '.join(group['items'])}")
    lines += ["", "## Work Experience"]
    for job in r["work_experience"]:
        lines.append(f"### {job['title']} | {job['date_range']}")
        line2 = job["company"] + (f" \u00b7 {job['team_context']}" if job.get("team_context") else "")
        lines.append(f"*{line2}*")
        for b in job["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.append("## Education")
    for ed in r["education"]:
        lines.append(f"- {ed['degree']}, {ed['institution']} ({ed['date']})")
    return strip_em_dashes("\n".join(lines))


def render_cover_letter_txt(cl: dict) -> str:
    return strip_em_dashes(cl["body"])


def render_cover_letter_md(cl: dict) -> str:
    return strip_em_dashes(cl["body"])


# --- DOCX --------------------------------------------------------------

ACCENT_RGB = RGBColor(0x1F, 0x38, 0x64)  # navy, sampled from the existing hand-written resumes


def _tight(paragraph, space_after=2, space_before=0):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    return paragraph


def render_resume_docx(r: dict, path: Path, body_pt: float = 10.5) -> None:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(body_pt)
    for section in doc.sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(r["name"])
    run.bold = True
    run.font.size = Pt(body_pt + 9)
    run.font.color.rgb = ACCENT_RGB
    _tight(name_p, space_after=2)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = contact_p.add_run(r["contact_line"])
    crun.font.size = Pt(body_pt - 1)
    _tight(contact_p, space_after=10)

    def add_heading(text):
        h = doc.add_paragraph()
        hr = h.add_run(text)
        hr.bold = True
        hr.font.size = Pt(body_pt + 1.5)
        hr.font.color.rgb = ACCENT_RGB
        _tight(h, space_after=1, space_before=10)
        pPr = h._p.get_or_add_pPr()
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        pBdr = pPr.makeelement(f"{ns}pBdr")
        bottom = pPr.makeelement(f"{ns}bottom")
        bottom.set(f"{ns}val", "single")
        bottom.set(f"{ns}sz", "6")
        bottom.set(f"{ns}color", "1F3864")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return h

    add_heading("SUMMARY")
    _tight(doc.add_paragraph(strip_em_dashes(r["summary"])), space_after=8)

    add_heading(r["skills_heading"])
    for group in r["skills"]:
        p = doc.add_paragraph()
        p.add_run(f"{group['category']}: ").bold = True
        p.add_run(strip_em_dashes(", ".join(group["items"])))
        _tight(p, space_after=2)

    add_heading("WORK EXPERIENCE")
    docx_usable_pt = 612 - 2 * 54  # page width minus the left/right margins set above (Pt(54) each)
    title_col, employer_col, date_col = compute_job_column_widths(r["work_experience"], body_pt, docx_usable_pt)
    for job in r["work_experience"]:
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        table.columns[0].width = Pt(title_col)
        table.columns[1].width = Pt(employer_col)
        table.columns[2].width = Pt(date_col)
        left, middle, right = table.rows[0].cells
        lp = left.paragraphs[0]
        lr = lp.add_run(job["title"])
        lr.bold = True
        _tight(lp, space_after=0)
        mp = middle.paragraphs[0]
        mr = mp.add_run(job["company"])
        mr.bold = True
        _tight(mp, space_after=0)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run(job["date_range"])
        rr.bold = True
        _tight(rp, space_after=0)

        if job.get("team_context"):
            cp = doc.add_paragraph()
            cr = cp.add_run(job["team_context"])
            cr.italic = True
            cr.font.color.rgb = ACCENT_RGB
            cr.font.size = Pt(body_pt - 0.5)
            _tight(cp, space_after=3)

        for i, b in enumerate(job["bullets"]):
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(strip_em_dashes(b))
            is_last = i == len(job["bullets"]) - 1
            _tight(bp, space_after=6 if is_last else 1)

    add_heading("EDUCATION")
    for ed in r["education"]:
        _tight(doc.add_paragraph(f"{ed['degree']}, {ed['institution']} ({ed['date']})"), space_after=1)

    doc.save(path)


def render_cover_letter_docx(cl: dict, path: Path) -> None:
    doc = Document()
    for para in strip_em_dashes(cl["body"]).split("\n\n"):
        doc.add_paragraph(para)
    doc.save(path)


# --- PDF (with page-fit: retries at smaller sizes until <=2 pages) -----

ACCENT_COLOR = colors.HexColor("#1F3864")  # sampled from the existing hand-written resumes

# Each tier: (body_pt, leading, bullet_space_after, top/bottom margin inches, name_pt, heading_pt)
PDF_TIERS = [
    (10.5, 13, 2, 0.55, 21, 12.5),
    (10, 12.5, 2, 0.5, 20, 12),
    (9.5, 12, 1.5, 0.45, 19, 11.5),
    (9, 11.5, 1, 0.4, 18, 11),
]


def _build_resume_story(r: dict, tier) -> list:
    body_pt, leading, bullet_space, _, name_pt, heading_pt = tier
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=body_pt, leading=leading, spaceAfter=3)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, spaceAfter=bullet_space, leftIndent=10)
    company_style = ParagraphStyle("Company", parent=body_style, textColor=ACCENT_COLOR, fontName="Helvetica-Oblique", spaceAfter=3)
    heading_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=heading_pt, textColor=ACCENT_COLOR,
                                    spaceBefore=heading_pt * 0.9, spaceAfter=2, fontName="Helvetica-Bold")
    name_style = ParagraphStyle("Name", parent=styles["Title"], fontSize=name_pt, textColor=ACCENT_COLOR, alignment=TA_CENTER, spaceAfter=2)
    contact_style = ParagraphStyle("Contact", parent=body_style, alignment=TA_CENTER, spaceAfter=10, fontSize=body_pt - 0.5)
    title_style = ParagraphStyle("JobTitle", parent=body_style, fontName="Helvetica-Bold", spaceAfter=0)
    employer_style = ParagraphStyle("JobEmployer", parent=body_style, fontName="Helvetica-Bold", spaceAfter=0)
    date_style = ParagraphStyle("JobDate", parent=title_style, alignment=2)

    def heading(text):
        return [Paragraph(text, heading_style), HRFlowable(width="100%", thickness=0.75, color=ACCENT_COLOR, spaceAfter=4)]

    story = [Paragraph(r["name"], name_style), Paragraph(r["contact_line"], contact_style)]
    story += heading("SUMMARY")
    story.append(Paragraph(strip_em_dashes(r["summary"]), body_style))
    story += heading(r["skills_heading"])
    for group in r["skills"]:
        story.append(Paragraph(f"<b>{group['category']}:</b> {strip_em_dashes(', '.join(group['items']))}", body_style))
    story += heading("WORK EXPERIENCE")
    usable_width = LETTER[0] - 2 * 0.7 * inch
    title_col, employer_col, date_col = compute_job_column_widths(r["work_experience"], body_pt, usable_width)
    for job in r["work_experience"]:
        row = Table(
            [[Paragraph(job["title"], title_style), Paragraph(job["company"], employer_style), Paragraph(job["date_range"], date_style)]],
            colWidths=[title_col, employer_col, date_col],
        )
        row.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(row)
        if job.get("team_context"):
            story.append(Paragraph(job["team_context"], company_style))
        for b in job["bullets"]:
            story.append(Paragraph(f"{BULLET_CHAR} {strip_em_dashes(b)}", bullet_style))
    story += heading("EDUCATION")
    for ed in r["education"]:
        story.append(Paragraph(f"{ed['degree']}, {ed['institution']} ({ed['date']})", body_style))
    return story


def render_resume_pdf(r: dict, path: Path, max_pages: int = 2):
    """Tries progressively smaller tiers until the rendered PDF fits within
    max_pages, or the smallest tier is reached. Returns (pages, body_pt)."""
    for tier in PDF_TIERS:
        _, _, _, margin_in, _, _ = tier
        doc = SimpleDocTemplate(
            str(path), pagesize=LETTER,
            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
            topMargin=margin_in * inch, bottomMargin=margin_in * inch,
        )
        doc.build(_build_resume_story(r, tier))
        if doc.page <= max_pages:
            return doc.page, tier[0]
    return doc.page, tier[0]


def render_cover_letter_pdf(cl: dict, path: Path) -> None:
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=11, leading=15, spaceAfter=10)
    doc = SimpleDocTemplate(str(path), pagesize=LETTER,
                             leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                             topMargin=0.9 * inch, bottomMargin=0.9 * inch)
    story = [Paragraph(p.replace("\n", "<br/>"), body_style) for p in strip_em_dashes(cl["body"]).split("\n\n")]
    doc.build(story)


# --- GitHub profile README (separate, template-driven LLM call) -------

README_REQUIRED_HEADERS = ["## \U0001f6e0\ufe0f Skills", "## \U0001f4bc Experience", "## \U0001f393 Education", "## \u2728 Career Highlights"]


def build_readme_context(kb: dict) -> dict:
    """Trimmed context for the README call -- same spirit as
    build_baseline_context, but variant-agnostic (the profile README isn't
    SDE- or SDET-specific) and includes the extra fields the resume schema
    doesn't carry: location, each education entry's field of study, and
    the career_narrative_notes differentiators list."""
    rules = kb["meta"]["output_rules"]
    skills = [
        {"category": CATEGORY_LABELS.get(k, k.replace("_", " ").title()), "items": v}
        for k, v in kb["skills"].items() if isinstance(v, list)
    ]
    work_experience = [
        {
            "title": job["title_by_variant"].get("SDE") or next(iter(job["title_by_variant"].values())),
            "company": job["company"],
            "team_context": job.get("team_context", ""),
            "date_range": f"{job['start_date']} - {job['end_date']}",
            "bullets": [b["text"] for b in job["bullets"] if not b["id"].endswith(("_alt", "_variant"))],
        }
        for job in kb["work_experience"]
    ]
    education = [
        {"degree": ed["degree"], "institution": ed["institution"],
         "field_of_study": ed.get("field of study", ""), "graduation_date": ed["graduation date"]}
        for ed in kb["education"]
    ]
    return {
        "output_rules": {"never_fabricate": rules["never_fabricate"], "never_use_em_dash": rules["never_use_em_dash"]},
        "personal_info": kb["personal_info"],
        "summary": kb["summary_variants"]["SDE"],
        "skills": skills,
        "work_experience": work_experience,
        "education": education,
        "career_highlights": kb.get("career_narrative_notes", {}).get("strongest_differentiators", []),
    }


def build_readme_system_prompt(kb: dict, template_text: str) -> str:
    context = build_readme_context(kb)
    return (
        "You are filling in a Markdown TEMPLATE for a GitHub profile README, "
        "using the candidate data below. Preserve the template's exact "
        "structure, Markdown syntax, emoji, and formatting (bold, headers, "
        "horizontal rules, bullet style) -- only replace the {{PLACEHOLDER}} "
        "tokens with real content. Follow the template's HTML-comment "
        "instructions for repeating blocks (one skills line per category, "
        "one Experience block per job, one Education block per entry, one "
        "bullet per career highlight). Reproduce career_highlights items "
        "VERBATIM -- do not reword, shorten, combine, or reorder them. Omit "
        "the '* team context *' line entirely for a job with no "
        "team_context. Follow output_rules exactly, especially "
        "never_fabricate and never_use_em_dash.\n\n"
        "=== TEMPLATE ===\n" + template_text + "\n\n"
        "=== CANDIDATE DATA (read for content only; do not include field "
        "names like team_context or graduation_date in your answer) ===\n"
        + json.dumps(context, indent=2)
        + "\n\n=== YOUR TASK ===\n"
        "Output ONLY the final, completed Markdown document -- no commentary, "
        "no markdown code fences around the whole thing, no leftover "
        "{{PLACEHOLDER}} tokens or HTML comments from the template. Start "
        "your response with the first line of the filled-in template."
    )


def extract_markdown(raw: str) -> str:
    """Strips a single wrapping ```...``` fence if the model put one around
    the whole document, despite being told not to."""
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n", "", text)
        text = re.sub(r"\n```$", "", text)
    return text.strip()


def validate_readme(md: str, expected_job_count: int) -> None:
    if not md.startswith("# "):
        raise ValueError("README does not start with a top-level '# ' heading.")
    missing = [h for h in README_REQUIRED_HEADERS if h not in md]
    if missing:
        raise ValueError(f"README is missing required section(s): {missing}")
    if "{{" in md or "}}" in md:
        raise ValueError("README still contains unfilled {{placeholder}} tokens.")
    if EM_DASH in md:
        raise ValueError("README contains an em dash, which violates never_use_em_dash.")
    job_headers = len(re.findall(r"^### ", md, flags=re.MULTILINE))
    if job_headers != expected_job_count:
        raise ValueError(f"README has {job_headers} job entries, expected {expected_job_count} (a job may have been dropped or duplicated).")


def call_llm_readme(kb: dict, template_text: str) -> str:
    base_url = os.environ.get("LITELLM_BASE_URL")
    api_key = os.environ.get("LITELLM_API_KEY")
    if not base_url or not api_key:
        print("LITELLM_BASE_URL and/or LITELLM_API_KEY are not set.", file=sys.stderr)
        sys.exit(1)

    client = openai.OpenAI(base_url=base_url.rstrip("/") + "/v1", api_key=api_key)
    system_prompt = build_readme_system_prompt(kb, template_text)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "Fill in the template now."},
    ]
    max_output_tokens = int(os.environ.get("LITELLM_MAX_TOKENS", "10000"))
    num_ctx = int(os.environ.get("OLLAMA_NUM_CTX", "16384"))
    expected_job_count = len(kb["work_experience"])

    est_input_tokens = len(system_prompt) // 4
    print(f"[README] Prompt size estimate: ~{est_input_tokens} input tokens.", file=sys.stderr)

    last_error = None
    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=180,
                extra_body={"options": {"num_ctx": num_ctx}}, messages=messages,
            )
        except openai.APIConnectionError as e:
            print(f"[README] Could not reach LiteLLM at {base_url}: {e}", file=sys.stderr)
            sys.exit(1)
        except openai.APIStatusError as e:
            print(f"[README] LiteLLM returned an error (HTTP {e.status_code}): {e.message}", file=sys.stderr)
            sys.exit(1)

        raw = response.choices[0].message.content or ""
        md = extract_markdown(raw)
        try:
            validate_readme(md, expected_job_count)
            return md
        except ValueError as e:
            last_error = e
            print(f"[README] Attempt {attempt + 1}: malformed output ({e}). Raw:\n{raw}\n", file=sys.stderr)
            messages.append({"role": "assistant", "content": raw})
            messages.append({
                "role": "user",
                "content": f"That response was invalid: {e}. Output ONLY the corrected, complete filled-in template, following all the same rules.",
            })

    print(f"[README] Model failed to produce a valid README after 2 attempts. Last error: {last_error}", file=sys.stderr)
    sys.exit(1)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", required=True, help="Path to resume_data.json")
    parser.add_argument("--out", required=True, help="Output directory (e.g. generated/)")
    parser.add_argument("--readme-out", default="README.md", help="Path for the GitHub profile README.md (default: repo root)")
    parser.add_argument("--readme-template", default="data/readme_template.md", help="Path to the README structure template")
    args = parser.parse_args()

    kb = json.loads(Path(args.data).read_text(encoding="utf-8"))
    out_dir = Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    for variant in VARIANTS:
        content = call_llm(kb, variant)
        r = content["resume"]
        cl = content["cover_letter"]

        base = out_dir / f"Dennis Dole Resume ({variant})"
        (base.with_suffix(".json")).write_text(json.dumps(r, indent=2), encoding="utf-8")
        (base.with_suffix(".txt")).write_text(render_resume_txt(r), encoding="utf-8")
        (base.with_suffix(".md")).write_text(render_resume_md(r), encoding="utf-8")

        pages, body_pt = render_resume_pdf(r, base.with_suffix(".pdf"))
        if pages > 2:
            print(f"WARNING: {variant} resume rendered at {pages} pages even at the smallest tier ({body_pt}pt).")
        else:
            print(f"{variant} resume: {pages} page(s) at {body_pt}pt.")
        render_resume_docx(r, base.with_suffix(".docx"), body_pt=body_pt)

        cl_base = out_dir / f"Dennis Dole Cover Letter ({variant})"
        (cl_base.with_suffix(".txt")).write_text(render_cover_letter_txt(cl), encoding="utf-8")
        render_cover_letter_docx(cl, cl_base.with_suffix(".docx"))
        render_cover_letter_pdf(cl, cl_base.with_suffix(".pdf"))

    # Separate LLM call, template-driven: fills data/readme_template.md
    # using resume_data.json, rather than reusing the resume call above.
    template_text = Path(args.readme_template).read_text(encoding="utf-8")
    readme_markdown = call_llm_readme(kb, template_text)
    readme_path = Path(args.readme_out)
    readme_path.write_text(readme_markdown, encoding="utf-8")
    print(f"Wrote GitHub profile README to {readme_path}")

    print(f"Wrote {len(VARIANTS)} resume variant(s) + 1 cover letter (5 formats each) to {out_dir}/")


if __name__ == "__main__":
    main()