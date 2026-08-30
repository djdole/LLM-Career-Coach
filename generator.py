#!/usr/bin/env python3
"""
Generates baseline (no-specific-JD) resumes (one per VARIANTS entry -
SDE and SDET by default), plus a cover
letter, from profile.json, in pdf/docx/txt/md/json formats.

Uses a self-hosted LiteLLM proxy (in front of Ollama, per that stack's
docker-compose.yml) rather than a paid hosted API, so this never spends
API credits and never fails due to account balance.

Tailoring a resume to a *specific* job posting is a different task (it
requires selecting/adapting content to a JD) and stays a separate,
manual, chat-based workflow - see profile.json's own
generation_workflow_for_llm for that path. This script does not do that.

Usage:
    python generate.py
    python generate.py --generate resume
    python generate.py --generate resume,cover_letter
    python generate.py --generate resume --generate readme
    python generate.py --analyze "paste JD text here"
    python generate.py --analyze path/to/job_posting.pdf
    python generate.py --analyze https://example.com/careers/some-job
    python generate.py --generate resume --analyze jd.txt

--generate controls WHAT gets built this run:
    * Omitted entirely: resumes, cover letters, and the README are all
      generated (the original, default behavior) - UNLESS --analyze was
      given and --generate was not, in which case nothing from --generate
      runs and this invocation does ONLY the analysis (see --analyze
      below). "profile" is never included in this default - it's
      opt-in only, see below.
    * Supplied with no value (e.g. a trailing `--generate` with nothing
      after it): nothing is generated.
    * Otherwise, its value is a comma-separated list of "resume",
      "cover_letter" (or "coverletter"), "readme", and/or "profile",
      and may be passed multiple times - the targets from every
      occurrence are combined.

--generate profile is a separate, opt-in workflow: rather than
generating resumes/cover letters/README FROM the knowledge base, it uses
source files dropped in the DATA folder (env var DATA, e.g. pdf/txt/json/
xml/docx documents) plus LiteLLM to build or non-destructively update a
profile.json. See generate_profile_draft() for the exact rules.

KNOWLEDGE_BASE (the knowledge base every target above reads from) is
usually a local path, but may instead be an http(s) URL - e.g. a raw
file URL into a private repo - so the knowledge base can be maintained
somewhere other than this checkout. See load_knowledge_base() and, for a
private source, the KNOWLEDGE_BASE_URL_TOKEN env var. --generate profile
still always writes its draft to a local path (KNOWLEDGE_BASE_DRAFT)
even when KNOWLEDGE_BASE is a URL - it has no way to push a draft back
to an arbitrary URL, so promoting the draft back to that source is a
manual step.

Resume/cover letter/README output (OUTPUT_FOLDER, README_OUTPUT)
normally lands in this same checkout, for the calling workflow/you to
commit. Setting OUTPUT_REPO decouples that: generated files are instead
written into a local clone of that OTHER git repo, which is committed
and pushed automatically at the end of the run - so this generator repo
and the repo that actually holds someone's checked-in resumes/cover
letters/README can be two different repos entirely. KNOWLEDGE_BASE_DRAFT
is unaffected - it's the source-of-truth knowledge base everything
else is generated FROM, not generated output itself, so it always
stays local. See sync_output_repo(), commit_and_push_output_repo(), and
the OUTPUT_REPO* env vars in .env.template.

--analyze is its own separate flag, independent of --generate: its value
IS the job description - literal JD text, a path to a local file
(pdf/docx/txt/md/json/xml), or a URL to fetch it from - see
resolve_job_description() for exactly how that value is interpreted. It
uses that plus data/profile.json and LiteLLM to estimate percentage
fit for that specific posting (0-100%), list the skills/qualifications
the JD calls for that aren't present in the knowledge base, and suggest
(preferably free) resources - tutorials, courses, books - to close each
gap. If the job description separates its qualifications into more than
one distinct list (e.g. "Required Qualifications" vs "Preferred
Qualifications"), a separate fit_percentage is produced per list instead
of one overall number - see the "fit_assessments" array in
ANALYSIS_PROMPT_TEMPLATE and validate_job_fit_analysis(). --analyze can
be combined with --generate in the same invocation to do both; see
call_llm_analyze_fit() for details. Its LLM prompt is loaded from
ANALYSIS_PROMPT_TEMPLATE (default: ANALYSIS_PROMPT.template.txt) and
filled in with the actual data by build_job_fit_prompt().
"""

import argparse
import datetime
import html
import json
import os
import re
import string
import subprocess
import sys
import types
import urllib.parse
from pathlib import Path

import openai
import httpx
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# The set of targets built whenever --generate is omitted entirely.
# "profile" is deliberately NOT a member of this set - it's a
# separate, opt-in maintenance workflow (see generate_profile_draft),
# not something that should run just because someone ran the script with
# no flags.
ALL_TARGETS = {"resume", "cover_letter", "readme"}

# Maps a normalized (lowercased, with "_"/"-" stripped) --generate token to
# its canonical target name. Both "cover_letter" and "coverletter" collapse
# to the same normalized key ("coverletter"), so either spelling works;
# likewise "profile" and "resumedata". Job-fit analysis is NOT one of
# these - it's triggered by the separate --analyze flag, not --generate
# (see build_arg_parser and main).
GENERATE_ALIASES = {
    "resume": "resume",
    "coverletter": "cover_letter",
    "readme": "readme",
    "profile": "profile",
    "resumedata": "profile",
}

# Every valid canonical target, default-generated or not - used for
# --generate's error message when an unknown value is supplied.
ALL_KNOWN_TARGETS = set(GENERATE_ALIASES.values())


def build_arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(
        description="Generates resumes, cover letters, and/or a GitHub "
        "profile README from profile.json (--generate), and/or runs "
        "a job-fit analysis against a job description (--analyze)."
    )
    parser.add_argument(
        "--generate",
        action="append",
        nargs="?",
        const="",
        default=None,
        metavar="resume,cover_letter,readme,profile",
        help="What to generate this run: a comma-separated list of "
        "resume, cover_letter, readme, and/or profile (may also be "
        "repeated, e.g. --generate resume --generate readme). Omit "
        "entirely to generate resume+cover_letter+readme (the default; "
        "profile is opt-in only and never included by default). "
        "Supply with no value to generate nothing. Unrelated to "
        "--analyze below, which is its own separate flag.",
    )
    parser.add_argument(
        "--analyze",
        default=None,
        metavar="JOB_DESCRIPTION",
        help="Run a job-fit analysis instead of/in addition to "
        "--generate. Its value IS the job description to evaluate fit "
        "against - literal text, a path to a local file (pdf/docx/txt/"
        "md/json/xml), or a URL to fetch it from. Uses that plus "
        "data/profile.json to estimate percentage fit (0-100), list "
        "missing skills/qualifications, and suggest (preferably free) "
        "resources to close each gap. Independent of --generate - pass "
        "both to do both in one run.",
    )
    return parser


def parse_generate_targets(raw_values: list) -> set:
    """
    Turns the raw list argparse collected for a repeatable, optional-value
    --generate flag into the set of canonical targets ("resume",
    "cover_letter", "readme") for this run.

    raw_values is one string per --generate occurrence: "" if that
    occurrence had no value, otherwise its (possibly comma-separated)
    value. Occurrences and comma-separated items within them are all
    unioned together, so e.g. ["resume", "cover_letter,readme"] and
    ["resume,cover_letter,readme"] are equivalent. An occurrence with no
    value contributes nothing, so a lone bare --generate (raw_values ==
    [""]) yields an empty set - callers should treat that as "generate
    nothing" rather than falling back to ALL_TARGETS.
    """
    targets = set()
    for raw in raw_values:
        for token in raw.split(","):
            normalized = token.strip().lower().replace("-", "").replace("_", "")
            if not normalized:
                continue
            canonical = GENERATE_ALIASES.get(normalized)
            if canonical is None:
                valid = ", ".join(sorted(ALL_KNOWN_TARGETS))
                print(
                    f"Unknown --generate value: {token.strip()!r}. Valid values: {valid}.",
                    file=sys.stderr,
                )
                sys.exit(1)
            targets.add(canonical)
    return targets

# Points at a self-hosted LiteLLM proxy instead of the Anthropic API. This
# stack runs LiteLLM in front of Ollama specifically as an OpenAI-compatible
# gateway (see docker-compose.yml), so we call LiteLLM directly rather than
# Open WebUI's own API layer. Avoids paid-API token usage entirely;
# generation runs against whatever local model Ollama has loaded.
#
# Required repo configuration (Settings -> Secrets and variables -> Actions):
#   Secrets:   LITELLM_BASE_URL   e.g. https://litellm.example.com
#                                 (points at the LiteLLM container's port,
#                                 ${LITELLM_PORT} in docker-compose.yml --
#                                 NOT Open WebUI's UI port)
#              LITELLM_API_KEY    same value as LITELLM_MASTER_KEY in that
#                                 stack's .env
#   Variables: LITELLM_MODEL      the model string LiteLLM proxies to, e.g.
#                                 "qwen3.6:latest" - NOT sensitive,
#                                 so it's a repo Variable rather than a
#                                 Secret.
#   Optional:  OLLAMA_NUM_CTX     context window size override (default
#                                 16384) - lower if your GPU can't hold
#                                 that much context for the model in use.
MODEL = os.environ.get("LITELLM_MODEL", "qwen3.6:latest")

# How long we wait for a *single* chat completion response before giving up
# on that attempt. There's no one correct default - it depends on your
# model size, hardware, LITELLM_MAX_TOKENS, and OLLAMA_NUM_CTX - so it's
# configurable via LITELLM_TIMEOUT rather than hardcoded. Keep it
# comfortably BELOW any reverse proxy's own read timeout in front of
# LiteLLM (proxy_read_timeout in nginx, etc.): if the proxy's timeout is
# shorter, IT kills the connection first, silently, and this timeout never
# gets the chance to produce the clearer error below. 550s is a safe
# default under nginx's commonly-recommended 600s.
LITELLM_TIMEOUT_SECONDS = float(os.environ.get("LITELLM_TIMEOUT", "550"))

# A single flat timeout number applies uniformly to the whole request,
# which isn't quite right for this use case: just reaching LiteLLM at all
# should fail fast (a few seconds) if it's unreachable, while actually
# *waiting on generated tokens* legitimately needs the full
# LITELLM_TIMEOUT_SECONDS budget for a local model generating up to
# LITELLM_MAX_TOKENS against OLLAMA_NUM_CTX of context. httpx.Timeout lets
# connect and read differ, so an unreachable/misconfigured proxy fails in
# ~10s with a clear connection error instead of burning the entire
# generation budget first only to fail for an unrelated reason. Passed to
# every client.chat.completions.create() call below in place of a bare
# `timeout=<seconds>`.
LLM_REQUEST_TIMEOUT = httpx.Timeout(
    connect=10.0,
    read=LITELLM_TIMEOUT_SECONDS,
    write=30.0,
    pool=LITELLM_TIMEOUT_SECONDS,
)

# How long Ollama keeps this model loaded in VRAM after a request, before
# unloading it. Left unset, a slow request (or a gap between the several
# calls a full run makes) risks the model unloading between calls, making
# the NEXT call pay a full model-load penalty on top of generation time.
# Explicit and generous relative to how long a full run takes end to end,
# so the model stays warm for the whole run rather than idling out mid-run.
# Ollama accepts a duration string ("30m", "1h") or seconds as a number.
LITELLM_KEEP_ALIVE = os.environ.get("LITELLM_KEEP_ALIVE", "30m")


def build_llm_client() -> openai.OpenAI:
    """
    Builds the single OpenAI-compatible client shared by every LLM call in
    a run. Reusing one client (and its underlying HTTP connection pool)
    instead of building a fresh one per call site avoids repeating the
    TCP/TLS handshake to LITELLM_BASE_URL for every one of the ~5-14
    sequential requests a full run makes.
    """
    base_url = os.environ.get("LITELLM_BASE_URL")
    api_key = os.environ.get("LITELLM_API_KEY")
    if not base_url or not api_key:
        print("LITELLM_BASE_URL and/or LITELLM_API_KEY are not set.", file=sys.stderr)
        sys.exit(1)
    return openai.OpenAI(base_url=base_url.rstrip("/") + "/v1", api_key=api_key)

# Maps the knowledge base's snake_case skill category keys to display
# labels matching the existing hand-written resumes' style, fed to the
# model as already-formatted so it only has to copy them, not invent
# formatting for them.
CATEGORY_LABELS = {
    "languages": "Languages",
    "apis_and_web_servers": "APIs & Web Servers",
    "test_automation_frameworks": "Test Automation Frameworks",
    "frontend_and_mobile_testing": "Frontend & Mobile Testing",
    "unit_integration_testing": "Unit & Integration Testing",
    "test_management_and_planning": "Test Management & Planning",
    "performance_testing": "Performance Testing",
    "ci_cd_and_devops": "CI/CD & DevOps",
    "databases_and_query_languages": "Databases & Query Languages",
    "version_control": "Version Control",
    "debugging_and_diagnostics": "Debugging & Diagnostics",
    "virtualization_and_infra": "Virtualization & Infra",
    "developer_tools_and_ides": "Developer Tools & IDEs",
    "ai_and_automation_tooling": "AI & Automation Tooling",
    "methodologies": "Methodologies",
    "collaboration_and_docs": "Collaboration & Docs",
    "monitoring_and_incident_mgmt": "Monitoring / Incident Mgmt",
}

SKILLS_HEADING_BY_VARIANT = {"SDE": "CORE TECHNICAL SKILLS", "SDET": "CORE SDET SKILLS"}
# Fallback heading for a VARIANTS entry (see load_file_location_settings)
# that isn't one of the two named above -- e.g. VARIANTS=SDE,SDET,SRE.
# Add an entry above for any variant that deserves custom wording
# instead of falling through to this generic one.
SKILLS_HEADING_FALLBACK = "CORE {variant} SKILLS"
BULLET_CHAR = "\u25cf"  # "●", matches the existing hand-written resumes' style
EM_DASH = "\u2014"


def compute_job_column_widths(work_experience: list, body_pt: float, total_pt: float, min_title_pt: float = 130) -> tuple:
    """
    Sizes the title/employer/date columns from the ACTUAL text in this
    resume at this font size, instead of fixed percentages - fixed splits
    silently overflow whenever a company name (e.g. 'Cosworth Tech Inc. /
    MAHLE Powertrain LLC') is longer than whatever guess produced the
    split, forcing it to wrap mid-name. Measured with reportlab's
    Helvetica-Bold metrics as a stand-in font for both PDF and DOCX --
    DOCX's actual font differs slightly, but widths are close enough that
    this still prevents wrapping in practice.

    Employer and date each get exactly what their longest actual value
    needs (plus a small buffer); title gets whatever's left, floored at
    min_title_pt so a long employer name can't crush the title column to
    nothing.
    """
    from reportlab.pdfbase.pdfmetrics import stringWidth
    pad = 8
    employer_pt = max(stringWidth(j["company"], "Helvetica-Bold", body_pt) for j in work_experience) + pad
    date_pt = max(stringWidth(j["date_range"], "Helvetica-Bold", body_pt) for j in work_experience) + pad
    title_pt = max(total_pt - employer_pt - date_pt, min_title_pt)
    return title_pt, employer_pt, date_pt


def strip_em_dashes(text: str) -> str:
    """Belt-and-suspenders fallback behind the never_use_em_dash rule the
    model is given - not a substitute for the model actually following it."""
    return text.replace(EM_DASH, ",")


# --- Prompt construction -------------------------------------------------

def build_baseline_context(kb: dict, variant: str) -> dict:
    """
    Trims the full knowledge base down to only what the BASELINE (no-JD)
    path needs for ONE variant (SDE or SDET), per generation_workflow_for_llm
    step 0's fallback.

    Trimming matters for two reasons: (1) token budget - the full KB is
    ~10k tokens alone, likely exceeding a local model's context window
    unless num_ctx is raised (see call_llm); (2) signal-to-noise - the
    full KB includes JD-tailoring-only fields (bullet variants, per-bullet
    themes/skills tags, the other variant's title, cover letter JD-specific
    building blocks) that have shown up verbatim in bad output before this
    trimming existed - fewer irrelevant JSON shapes nearby means less for
    a smaller model to latch onto instead of the requested schema.

    Skill category labels are pre-formatted here (see CATEGORY_LABELS) so
    the model only has to copy them, not invent formatting.
    """
    rules = kb["meta"]["output_rules"]
    summary = kb["summary_variants"].get(variant) or kb["summary_variants"]["SDE"]
    skills = [
        {"category": CATEGORY_LABELS.get(k, k.replace("_", " ").title()), "items": v}
        for k, v in kb["skills"].items() if isinstance(v, list)
    ]
    work_experience = [
        {
            "title": job["title_by_variant"].get(variant) or next(iter(job["title_by_variant"].values())),
            "company": job["company"],
            "team_context": job.get("team_context", ""),
            "date_range": f"{job['start_date']} - {job['end_date']}",
            "bullets": [b["text"] for b in job["bullets"] if not b["id"].endswith(("_alt", "_variant"))],
        }
        for job in kb["work_experience"]
    ]
    # The raw KB uses the key "graduation date" (with a space); the output
    # schema calls it "date". Renaming it here - rather than relying on the
    # model to perform that rename - is what actually fixed education[]
    # entries coming back without a date at all.
    education = [
        {"degree": ed["degree"], "institution": ed["institution"], "date": ed["graduation date"]}
        for ed in kb["education"]
    ]
    return {
        "output_rules": {
            "never_fabricate": rules["never_fabricate"],
            "never_use_em_dash": rules["never_use_em_dash"],
        },
        "personal_info": kb["personal_info"],
        "education": education,
        "summary": summary,
        "skills_heading": SKILLS_HEADING_BY_VARIANT.get(variant, SKILLS_HEADING_FALLBACK.format(variant=variant.upper())),
        "skills": skills,
        "work_experience": work_experience,
        "cover_letter_generic_template": kb["cover_letter_building_blocks"]["generic_fallback_template"],
    }


def build_resume_fill_prompt(kb: dict, variant: str, template_text: str) -> str:
    """Fills RESUME_TEMPLATE using the trimmed baseline context - same
    spirit as build_readme_system_prompt, adapted for the resume's
    pipe-delimited, code-parsed structure."""
    context = build_baseline_context(kb, variant)
    return (
        f"You are filling in a plain-text resume TEMPLATE for a BASELINE "
        f"{variant} resume (no specific job description provided), using "
        "the candidate data below. Preserve the template's exact structure "
        "and formatting - section header text (SUMMARY, WORK EXPERIENCE, "
        "EDUCATION), the bullet character, and especially the exact ' | ' "
        "(space-pipe-space) delimiters on the job-header and education "
        "lines, since those are parsed by code afterward and must be exact. "
        "Only replace the {{PLACEHOLDER}} tokens with real content. Follow "
        "the template's HTML-comment instructions for repeating blocks (one "
        "skills line per category, one Experience block per job, one "
        "Education line per entry). Omit the team-context line entirely "
        "for a job with no team_context. Follow output_rules exactly, "
        "especially never_fabricate and never_use_em_dash.\n\n"
        "=== TEMPLATE ===\n" + template_text + "\n\n"
        "=== CANDIDATE DATA (read for content only; do not include field "
        "names like team_context or date_range in your answer) ===\n"
        + json.dumps(context, indent=2)
        + "\n\n=== YOUR TASK ===\n"
        "Output ONLY the final, completed document - no commentary, no "
        "markdown code fences around the whole thing, no leftover "
        "{{PLACEHOLDER}} tokens or HTML comments from the template."
    )


def parse_filled_resume(text: str) -> dict:
    """Parses a filled RESUME_TEMPLATE (see build_resume_fill_prompt) back
    into the structured dict the existing renderers expect. Raises
    ValueError on any structural mismatch, which the caller uses to
    trigger a corrective retry rather than crash deep inside rendering."""
    lines = text.strip("\n").split("\n")
    i = 0

    def skip_blank():
        nonlocal i
        while i < len(lines) and lines[i].strip() == "":
            i += 1

    if len(lines) < 2:
        raise ValueError("Output too short to contain name/contact lines.")
    name = lines[i].strip(); i += 1
    contact_line = lines[i].strip(); i += 1

    skip_blank()
    if i >= len(lines) or lines[i].strip() != "SUMMARY":
        raise ValueError(f"Expected 'SUMMARY' header, got: {lines[i].strip() if i < len(lines) else 'EOF'!r}")
    i += 1
    skip_blank()
    summary_lines = []
    while i < len(lines) and lines[i].strip() != "":
        summary_lines.append(lines[i].strip())
        i += 1
    summary = " ".join(summary_lines)

    skip_blank()
    if i >= len(lines):
        raise ValueError("Output ended before a skills_heading line.")
    skills_heading = lines[i].strip(); i += 1
    skills = []
    while i < len(lines) and lines[i].strip() and lines[i].strip() != "WORK EXPERIENCE":
        m = re.match(r"^(.+?):\s*(.+)$", lines[i].strip())
        if not m:
            raise ValueError(f"Could not parse skills line: {lines[i]!r}")
        category, items_str = m.groups()
        skills.append({"category": category, "items": [x.strip() for x in items_str.split(",") if x.strip()]})
        i += 1
    skip_blank()
    if i >= len(lines) or lines[i].strip() != "WORK EXPERIENCE":
        raise ValueError(f"Expected 'WORK EXPERIENCE' header, got: {lines[i].strip() if i < len(lines) else 'EOF'!r}")
    i += 1
    skip_blank()

    work_experience = []
    while i < len(lines) and lines[i].strip() and lines[i].strip() != "EDUCATION":
        header = lines[i].strip(); i += 1
        parts = [p.strip() for p in header.split("|")]
        if len(parts) != 3:
            raise ValueError(f"Job header line not in 'Title | Company | Dates' shape: {header!r}")
        title, company, date_range = parts
        team_context = ""
        if i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(BULLET_CHAR):
            team_context = lines[i].strip()
            i += 1
        bullets = []
        while i < len(lines) and lines[i].strip().startswith(BULLET_CHAR):
            bullets.append(lines[i].strip()[1:].strip())
            i += 1
        if not bullets:
            raise ValueError(f"Job '{title}' has no bullets.")
        work_experience.append({"title": title, "company": company, "date_range": date_range,
                                 "team_context": team_context, "bullets": bullets})
        skip_blank()

    if i >= len(lines) or lines[i].strip() != "EDUCATION":
        raise ValueError(f"Expected 'EDUCATION' header, got: {lines[i].strip() if i < len(lines) else 'EOF'!r}")
    i += 1
    skip_blank()
    education = []
    while i < len(lines) and lines[i].strip():
        parts = [p.strip() for p in lines[i].split("|")]
        if len(parts) != 3:
            raise ValueError(f"Education line not in 'Degree | Institution | Date' shape: {lines[i]!r}")
        degree, institution, date = parts
        education.append({"degree": degree, "institution": institution, "date": date})
        i += 1

    return {"name": name, "contact_line": contact_line, "skills_heading": skills_heading,
            "summary": summary, "skills": skills, "work_experience": work_experience, "education": education}


def call_llm_fill_resume(client: openai.OpenAI, kb: dict, variant: str, template_text: str) -> dict:
    system_prompt = build_resume_fill_prompt(kb, variant, template_text)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "Fill in the template now."},
    ]
    max_output_tokens = int(os.environ.get("LITELLM_MAX_TOKENS", "10000"))
    num_ctx = int(os.environ.get("OLLAMA_NUM_CTX", "16384"))
    expected_job_count = len(kb["work_experience"])

    est_input_tokens = len(system_prompt) // 4
    print(f"[{variant}] Resume prompt size estimate: ~{est_input_tokens} input tokens.", file=sys.stderr)

    last_error = None
    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=LLM_REQUEST_TIMEOUT,
                extra_body={"options": {"num_ctx": num_ctx}, "keep_alive": LITELLM_KEEP_ALIVE}, messages=messages,
            )
        except openai.APIConnectionError as e:
            print(f"[{variant}] Could not reach LiteLLM at {client.base_url}: {e}", file=sys.stderr)
            sys.exit(1)
        except openai.APIStatusError as e:
            print(f"[{variant}] LiteLLM returned an error (HTTP {e.status_code}): {e.message}", file=sys.stderr)
            sys.exit(1)

        raw = response.choices[0].message.content or ""
        text = extract_markdown(raw)
        try:
            parsed = parse_filled_resume(text)
            if len(parsed["work_experience"]) != expected_job_count:
                raise ValueError(f"Got {len(parsed['work_experience'])} jobs, expected {expected_job_count}.")
            return parsed
        except ValueError as e:
            last_error = e
            print(f"[{variant}] Attempt {attempt + 1}: malformed resume output ({e}). Raw:\n{raw}\n", file=sys.stderr)
            messages.append({"role": "assistant", "content": raw})
            messages.append({
                "role": "user",
                "content": f"That response was invalid: {e}. Output ONLY the corrected, complete filled-in template, following all the same rules.",
            })

    print(f"[{variant}] Model failed to produce a valid filled resume after 2 attempts. Last error: {last_error}", file=sys.stderr)
    sys.exit(1)


REQUIRED_COVER_LETTER_KEYS = {"body"}


def extract_json_object(raw: str) -> str:
    """Pulls the first balanced {...} block out of a string that may
    contain markdown fences and/or prose before/after it."""
    start = raw.find("{")
    if start == -1:
        raise ValueError("No '{' found in model output.")
    depth = 0
    for i, ch in enumerate(raw[start:], start=start):
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return raw[start:i + 1]
    raise ValueError("No balanced closing '}' found in model output.")


def build_cover_letter_prompt(kb: dict, variant: str) -> str:
    context = build_baseline_context(kb, variant)
    return (
        "You are producing a BASELINE cover letter (no specific job "
        "description provided), for the candidate described below. Use "
        "cover_letter_generic_template as the cover letter (lightly adapt "
        "it, keep 'Dear Hiring Manager'). Follow output_rules exactly, "
        "especially never_fabricate and never_use_em_dash.\n\n"
        "=== CANDIDATE DATA ===\n" + json.dumps({
            "output_rules": context["output_rules"],
            "cover_letter_generic_template": context["cover_letter_generic_template"],
        }, indent=2)
        + "\n\n=== YOUR TASK ===\n"
        'Respond with ONLY a JSON object of the shape {"body": "full cover '
        'letter text, paragraphs separated by a blank line"}. No prose '
        "before or after it, no markdown code fences. Your entire response "
        "must start with { and end with }."
    )


def call_llm_cover_letter(client: openai.OpenAI, kb: dict, variant: str) -> dict:
    system_prompt = build_cover_letter_prompt(kb, variant)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "Generate the cover letter JSON now."},
    ]
    max_output_tokens = int(os.environ.get("LITELLM_MAX_TOKENS", "10000"))
    num_ctx = int(os.environ.get("OLLAMA_NUM_CTX", "16384"))

    last_error = None
    for attempt in range(2):
        try:
            try:
                response = client.chat.completions.create(
                    model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=LLM_REQUEST_TIMEOUT,
                    response_format={"type": "json_object"},
                    extra_body={"options": {"num_ctx": num_ctx}, "keep_alive": LITELLM_KEEP_ALIVE}, messages=messages,
                )
            except openai.BadRequestError:
                response = client.chat.completions.create(
                    model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=LLM_REQUEST_TIMEOUT,
                    extra_body={"options": {"num_ctx": num_ctx}, "keep_alive": LITELLM_KEEP_ALIVE}, messages=messages,
                )
        except openai.APIConnectionError as e:
            print(f"[{variant}] Could not reach LiteLLM at {client.base_url}: {e}", file=sys.stderr)
            sys.exit(1)
        except openai.APIStatusError as e:
            print(f"[{variant}] LiteLLM returned an error (HTTP {e.status_code}): {e.message}", file=sys.stderr)
            sys.exit(1)

        raw = response.choices[0].message.content or ""
        try:
            cl = json.loads(extract_json_object(raw))
            if REQUIRED_COVER_LETTER_KEYS - set(cl.keys()):
                raise ValueError("cover letter is missing required key: 'body'")
            return cl
        except (ValueError, json.JSONDecodeError) as e:
            last_error = e
            print(f"[{variant}] Attempt {attempt + 1}: malformed cover letter output ({e}). Raw:\n{raw}\n", file=sys.stderr)
            messages.append({"role": "assistant", "content": raw})
            messages.append({
                "role": "user",
                "content": f"That response was invalid: {e}. Respond again with ONLY the JSON object described earlier.",
            })

    print(f"[{variant}] Model failed to produce a valid cover letter after 2 attempts. Last error: {last_error}", file=sys.stderr)
    sys.exit(1)


# --- Plain text / Markdown -------------------------------------------------

def render_resume_txt(r: dict) -> str:
    lines = [r["name"], r["contact_line"], "", "SUMMARY", r["summary"], "", r["skills_heading"]]
    for group in r["skills"]:
        lines.append(f"{group['category']}: {', '.join(group['items'])}")
    lines += ["", "WORK EXPERIENCE"]
    for job in r["work_experience"]:
        lines.append(f"{job['title']} {job['date_range']}")
        line2 = job["company"] + (f" \u00b7 {job['team_context']}" if job.get("team_context") else "")
        lines.append(line2)
        for b in job["bullets"]:
            lines.append(f"{BULLET_CHAR} {b}")
        lines.append("")
    lines.append("EDUCATION")
    for ed in r["education"]:
        lines.append(f"{ed['degree']}")
        lines.append(f"{ed['institution']} ({ed['date']})")
    return strip_em_dashes("\n".join(lines))


def render_resume_md(r: dict) -> str:
    lines = [f"# {r['name']}", r["contact_line"], "", "## Summary", r["summary"], "", f"## {r['skills_heading'].title()}"]
    for group in r["skills"]:
        lines.append(f"- **{group['category']}:** {', '.join(group['items'])}")
    lines += ["", "## Work Experience"]
    for job in r["work_experience"]:
        lines.append(f"### {job['title']} | {job['date_range']}")
        line2 = job["company"] + (f" \u00b7 {job['team_context']}" if job.get("team_context") else "")
        lines.append(f"*{line2}*")
        for b in job["bullets"]:
            lines.append(f"- {b}")
        lines.append("")
    lines.append("## Education")
    for ed in r["education"]:
        lines.append(f"- {ed['degree']}, {ed['institution']} ({ed['date']})")
    return strip_em_dashes("\n".join(lines))


def render_cover_letter_txt(cl: dict) -> str:
    return strip_em_dashes(cl["body"])


def render_cover_letter_md(cl: dict) -> str:
    return strip_em_dashes(cl["body"])


# --- DOCX --------------------------------------------------------------

ACCENT_RGB = RGBColor(0x1F, 0x38, 0x64)  # navy, sampled from the existing hand-written resumes


def _tight(paragraph, space_after=2, space_before=0):
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.space_before = Pt(space_before)
    return paragraph


def render_resume_docx(r: dict, path: Path, body_pt: float = 10.5) -> None:
    doc = Document()
    doc.styles["Normal"].font.size = Pt(body_pt)
    for section in doc.sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(r["name"])
    run.bold = True
    run.font.size = Pt(body_pt + 9)
    run.font.color.rgb = ACCENT_RGB
    _tight(name_p, space_after=2)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    crun = contact_p.add_run(r["contact_line"])
    crun.font.size = Pt(body_pt - 1)
    _tight(contact_p, space_after=10)

    def add_heading(text):
        h = doc.add_paragraph()
        hr = h.add_run(text)
        hr.bold = True
        hr.font.size = Pt(body_pt + 1.5)
        hr.font.color.rgb = ACCENT_RGB
        _tight(h, space_after=1, space_before=10)
        pPr = h._p.get_or_add_pPr()
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        pBdr = pPr.makeelement(f"{ns}pBdr")
        bottom = pPr.makeelement(f"{ns}bottom")
        bottom.set(f"{ns}val", "single")
        bottom.set(f"{ns}sz", "6")
        bottom.set(f"{ns}color", "1F3864")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return h

    add_heading("SUMMARY")
    _tight(doc.add_paragraph(strip_em_dashes(r["summary"])), space_after=8)

    add_heading(r["skills_heading"])
    for group in r["skills"]:
        p = doc.add_paragraph()
        p.add_run(f"{group['category']}: ").bold = True
        p.add_run(strip_em_dashes(", ".join(group["items"])))
        _tight(p, space_after=2)

    add_heading("WORK EXPERIENCE")
    docx_usable_pt = 612 - 2 * 54  # page width minus the left/right margins set above (Pt(54) each)
    title_col, employer_col, date_col = compute_job_column_widths(r["work_experience"], body_pt, docx_usable_pt)
    for job in r["work_experience"]:
        table = doc.add_table(rows=1, cols=3)
        table.autofit = True
        table.columns[0].width = Pt(title_col)
        table.columns[1].width = Pt(employer_col)
        table.columns[2].width = Pt(date_col)
        left, middle, right = table.rows[0].cells
        lp = left.paragraphs[0]
        lr = lp.add_run(job["title"])
        lr.bold = True
        _tight(lp, space_after=0)
        mp = middle.paragraphs[0]
        mr = mp.add_run(job["company"])
        mr.bold = True
        _tight(mp, space_after=0)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = rp.add_run(job["date_range"])
        rr.bold = True
        _tight(rp, space_after=0)

        if job.get("team_context"):
            cp = doc.add_paragraph()
            cr = cp.add_run(job["team_context"])
            cr.italic = True
            cr.font.color.rgb = ACCENT_RGB
            cr.font.size = Pt(body_pt - 0.5)
            _tight(cp, space_after=3)

        for i, b in enumerate(job["bullets"]):
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(strip_em_dashes(b))
            is_last = i == len(job["bullets"]) - 1
            _tight(bp, space_after=6 if is_last else 1)

    add_heading("EDUCATION")
    for ed in r["education"]:
        _tight(doc.add_paragraph(f"{ed['degree']}, {ed['institution']} ({ed['date']})"), space_after=1)

    doc.save(path)


def render_cover_letter_docx(cl: dict, path: Path) -> None:
    doc = Document()
    for para in strip_em_dashes(cl["body"]).split("\n\n"):
        doc.add_paragraph(para)
    doc.save(path)


# --- PDF (with page-fit: retries at smaller sizes until <=2 pages) -----

ACCENT_COLOR = colors.HexColor("#1F3864")  # sampled from the existing hand-written resumes

# Each tier: (body_pt, leading, bullet_space_after, top/bottom margin inches, name_pt, heading_pt)
PDF_TIERS = [
    (10.5, 13, 2, 0.55, 21, 12.5),
    (10, 12.5, 2, 0.5, 20, 12),
    (9.5, 12, 1.5, 0.45, 19, 11.5),
    (9, 11.5, 1, 0.4, 18, 11),
]


def _build_resume_story(r: dict, tier) -> list:
    body_pt, leading, bullet_space, _, name_pt, heading_pt = tier
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=body_pt, leading=leading, spaceAfter=3)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, spaceAfter=bullet_space, leftIndent=10)
    company_style = ParagraphStyle("Company", parent=body_style, textColor=ACCENT_COLOR, fontName="Helvetica-Oblique", spaceAfter=3)
    heading_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=heading_pt, textColor=ACCENT_COLOR,
                                    spaceBefore=heading_pt * 0.9, spaceAfter=2, fontName="Helvetica-Bold")
    name_style = ParagraphStyle("Name", parent=styles["Title"], fontSize=name_pt, textColor=ACCENT_COLOR, alignment=TA_CENTER, spaceAfter=2)
    contact_style = ParagraphStyle("Contact", parent=body_style, alignment=TA_CENTER, spaceAfter=10, fontSize=body_pt - 0.5)
    title_style = ParagraphStyle("JobTitle", parent=body_style, fontName="Helvetica-Bold", spaceAfter=0)
    employer_style = ParagraphStyle("JobEmployer", parent=body_style, fontName="Helvetica-Bold", spaceAfter=0)
    date_style = ParagraphStyle("JobDate", parent=title_style, alignment=2)

    def heading(text):
        return [Paragraph(text, heading_style), HRFlowable(width="100%", thickness=0.75, color=ACCENT_COLOR, spaceAfter=4)]

    story = [Paragraph(r["name"], name_style), Paragraph(r["contact_line"], contact_style)]
    story += heading("SUMMARY")
    story.append(Paragraph(strip_em_dashes(r["summary"]), body_style))
    story += heading(r["skills_heading"])
    for group in r["skills"]:
        story.append(Paragraph(f"<b>{group['category']}:</b> {strip_em_dashes(', '.join(group['items']))}", body_style))
    story += heading("WORK EXPERIENCE")
    usable_width = LETTER[0] - 2 * 0.7 * inch
    title_col, employer_col, date_col = compute_job_column_widths(r["work_experience"], body_pt, usable_width)
    for job in r["work_experience"]:
        row = Table(
            [[Paragraph(job["title"], title_style), Paragraph(job["company"], employer_style), Paragraph(job["date_range"], date_style)]],
            colWidths=[title_col, employer_col, date_col],
        )
        row.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0), ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0), ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(row)
        if job.get("team_context"):
            story.append(Paragraph(job["team_context"], company_style))
        for b in job["bullets"]:
            story.append(Paragraph(f"{BULLET_CHAR} {strip_em_dashes(b)}", bullet_style))
    story += heading("EDUCATION")
    for ed in r["education"]:
        story.append(Paragraph(f"{ed['degree']}, {ed['institution']} ({ed['date']})", body_style))
    return story


def render_resume_pdf(r: dict, path: Path, max_pages: int = 2):
    """Tries progressively smaller tiers until the rendered PDF fits within
    max_pages, or the smallest tier is reached. Returns (pages, body_pt)."""
    for tier in PDF_TIERS:
        _, _, _, margin_in, _, _ = tier
        doc = SimpleDocTemplate(
            str(path), pagesize=LETTER,
            leftMargin=0.7 * inch, rightMargin=0.7 * inch,
            topMargin=margin_in * inch, bottomMargin=margin_in * inch,
        )
        doc.build(_build_resume_story(r, tier))
        if doc.page <= max_pages:
            return doc.page, tier[0]
    return doc.page, tier[0]


def render_cover_letter_pdf(cl: dict, path: Path) -> None:
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=11, leading=15, spaceAfter=10)
    doc = SimpleDocTemplate(str(path), pagesize=LETTER,
                             leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                             topMargin=0.9 * inch, bottomMargin=0.9 * inch)
    story = [Paragraph(p.replace("\n", "<br/>"), body_style) for p in strip_em_dashes(cl["body"]).split("\n\n")]
    doc.build(story)


# --- GitHub profile README (separate, template-driven LLM call) -------

README_REQUIRED_HEADERS = ["## \U0001f6e0\ufe0f Skills", "## \U0001f4bc Experience", "## \U0001f393 Education", "## \u2728 Career Highlights"]


def build_tagged_email(email: str, tag_address: str) -> str:
    """
    Builds the email address used for the README's mailto: link, applying
    "plus addressing" (RFC 5233 subaddressing, e.g. Gmail/Outlook/etc.)
    if EMAIL_TAG_ADDRESS is set: 'jane@example.com' + 'resume' becomes
    'jane+resume@example.com', so the visible/displayed email can stay
    exactly as-is while replies to (or the sender seeing) the mailto
    link's address reveal it came from the README specifically.

    Returns email UNCHANGED (no '+' inserted) if tag_address is blank,
    whitespace-only, or unset -- and also if email doesn't look like a
    single-@ address, since a malformed knowledge-base email shouldn't
    crash the whole README run over a cosmetic feature.
    """
    tag_address = (tag_address or "").strip()
    if not tag_address:
        return email
    local, sep, domain = email.partition("@")
    if not sep:
        return email
    return f"{local}+{tag_address}@{domain}"


def build_readme_context(kb: dict) -> dict:
    """Trimmed context for the README call - same spirit as
    build_baseline_context, but variant-agnostic (the profile README isn't
    SDE- or SDET-specific) and includes the extra fields the resume schema
    doesn't carry: location, each education entry's field of study, and
    the career_narrative_notes differentiators list.

    personal_info also gets an added email_mailto field (see
    build_tagged_email) alongside the knowledge base's own untouched
    email -- the template uses email for the DISPLAYED text and
    email_mailto for the mailto: link's address, so an optional
    EMAIL_TAG_ADDRESS can be embedded in the link the reader clicks
    without changing the email address shown on the page."""
    rules = kb["meta"]["output_rules"]
    skills = [
        {"category": CATEGORY_LABELS.get(k, k.replace("_", " ").title()), "items": v}
        for k, v in kb["skills"].items() if isinstance(v, list)
    ]
    work_experience = [
        {
            "title": job["title_by_variant"].get("SDE") or next(iter(job["title_by_variant"].values())),
            "company": job["company"],
            "team_context": job.get("team_context", ""),
            "date_range": f"{job['start_date']} - {job['end_date']}",
            "bullets": [b["text"] for b in job["bullets"] if not b["id"].endswith(("_alt", "_variant"))],
        }
        for job in kb["work_experience"]
    ]
    education = [
        {"degree": ed["degree"], "institution": ed["institution"],
         "field_of_study": ed.get("field of study", ""), "graduation_date": ed["graduation date"]}
        for ed in kb["education"]
    ]
    personal_info = dict(kb["personal_info"])
    personal_info["email_mailto"] = build_tagged_email(
        kb["personal_info"]["email"], os.environ.get("EMAIL_TAG_ADDRESS", "")
    )
    return {
        "output_rules": {"never_fabricate": rules["never_fabricate"], "never_use_em_dash": rules["never_use_em_dash"]},
        "personal_info": personal_info,
        "summary": kb["summary_variants"]["SDE"],
        "skills": skills,
        "work_experience": work_experience,
        "education": education,
        "career_highlights": kb.get("career_narrative_notes", {}).get("strongest_differentiators", []),
    }


def build_readme_system_prompt(kb: dict, template_text: str) -> str:
    context = build_readme_context(kb)
    return (
        "You are filling in a Markdown TEMPLATE for a GitHub profile README, "
        "using the candidate data below. Preserve the template's exact "
        "structure, Markdown syntax, emoji, and formatting (bold, headers, "
        "horizontal rules, bullet style) - only replace the {{PLACEHOLDER}} "
        "tokens with real content. Follow the template's HTML-comment "
        "instructions for repeating blocks (one skills line per category, "
        "one Experience block per job, one Education block per entry, one "
        "bullet per career highlight). Reproduce career_highlights items "
        "VERBATIM - do not reword, shorten, combine, or reorder them. Omit "
        "the '* team context *' line entirely for a job with no "
        "team_context. Follow output_rules exactly, especially "
        "never_fabricate and never_use_em_dash.\n\n"
        "IMPORTANT - two distinct email fields: personal_info.email is "
        "the DISPLAYED email text and personal_info.email_mailto is the "
        "address inside the mailto: link. They may differ (email_mailto "
        "can carry an extra '+tag' for the link's address only) or be "
        "identical - either way, use email EXACTLY where the template "
        "shows visible email text, and email_mailto EXACTLY inside every "
        "mailto: URL. Never swap them, and never merge them into one "
        "value.\n\n"
        "=== TEMPLATE ===\n" + template_text + "\n\n"
        "=== CANDIDATE DATA (read for content only; do not include field "
        "names like team_context or graduation_date in your answer) ===\n"
        + json.dumps(context, indent=2)
        + "\n\n=== YOUR TASK ===\n"
        "Output ONLY the final, completed Markdown document - no commentary, "
        "no markdown code fences around the whole thing, no leftover "
        "{{PLACEHOLDER}} tokens or HTML comments from the template. Start "
        "your response with the first line of the filled-in template."
    )


def extract_markdown(raw: str) -> str:
    """Strips a single wrapping ```...``` fence if the model put one around
    the whole document, despite being told not to."""
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n", "", text)
        text = re.sub(r"\n```$", "", text)
    return text.strip()


def validate_readme(md: str, expected_job_count: int) -> None:
    if not re.search(r"^# ", md, flags=re.MULTILINE):
        raise ValueError("README is missing its top-level '# ' heading.")
    missing = [h for h in README_REQUIRED_HEADERS if h not in md]
    if missing:
        raise ValueError(f"README is missing required section(s): {missing}")
    if "{{" in md or "}}" in md:
        raise ValueError("README still contains unfilled {{placeholder}} tokens.")
    if EM_DASH in md:
        raise ValueError("README contains an em dash, which violates never_use_em_dash.")
    job_headers = len(re.findall(r"^### ", md, flags=re.MULTILINE))
    if job_headers != expected_job_count:
        raise ValueError(f"README has {job_headers} job entries, expected {expected_job_count} (a job may have been dropped or duplicated).")


def call_llm_readme(client: openai.OpenAI, kb: dict, template_text: str) -> str:
    system_prompt = build_readme_system_prompt(kb, template_text)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "Fill in the template now."},
    ]
    max_output_tokens = int(os.environ.get("LITELLM_MAX_TOKENS", "10000"))
    num_ctx = int(os.environ.get("OLLAMA_NUM_CTX", "16384"))
    expected_job_count = len(kb["work_experience"])

    est_input_tokens = len(system_prompt) // 4
    print(f"[README] Prompt size estimate: ~{est_input_tokens} input tokens.", file=sys.stderr)

    last_error = None
    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=LLM_REQUEST_TIMEOUT,
                extra_body={"options": {"num_ctx": num_ctx}, "keep_alive": LITELLM_KEEP_ALIVE}, messages=messages,
            )
        except openai.APIConnectionError as e:
            print(f"[README] Could not reach LiteLLM at {client.base_url}: {e}", file=sys.stderr)
            sys.exit(1)
        except openai.APIStatusError as e:
            print(f"[README] LiteLLM returned an error (HTTP {e.status_code}): {e.message}", file=sys.stderr)
            sys.exit(1)

        raw = response.choices[0].message.content or ""
        md = extract_markdown(raw)
        try:
            validate_readme(md, expected_job_count)
            return md
        except ValueError as e:
            last_error = e
            print(f"[README] Attempt {attempt + 1}: malformed output ({e}). Raw:\n{raw}\n", file=sys.stderr)
            messages.append({"role": "assistant", "content": raw})
            messages.append({
                "role": "user",
                "content": f"That response was invalid: {e}. Output ONLY the corrected, complete filled-in template, following all the same rules.",
            })

    print(f"[README] Model failed to produce a valid README after 2 attempts. Last error: {last_error}", file=sys.stderr)
    sys.exit(1)


def extract_text_from_source_file(path: Path) -> str:
    """
    Best-effort plain-text extraction from a DATA-folder source file, for
    feeding to the LLM in generate_profile_draft(). Supports the
    formats a resume-adjacent document is likely to show up in: json,
    txt/md, xml, docx, pdf. Unsupported or unreadable files return "" (and
    are logged), rather than raising, so one bad file in the DATA folder
    doesn't abort the whole run.
    """
    suffix = path.suffix.lower()
    try:
        if suffix in (".json", ".txt", ".md", ".xml"):
            return path.read_text(encoding="utf-8")
        if suffix == ".docx":
            doc = Document(str(path))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        if suffix == ".pdf":
            reader = PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        print(f"[profile] Could not read {path.name}: {e}", file=sys.stderr)
        return ""

    print(f"[profile] Skipping unsupported file type: {path.name}", file=sys.stderr)
    return ""


def build_source_file_list(data_dir: Path, knowledge_base_path: Path | None, draft_path: Path) -> list:
    """
    Lists the candidate source files sitting in the DATA folder: every
    regular, non-hidden file EXCEPT the knowledge base file itself (which
    may well live in the same folder, e.g. data/profile.json) and any
    pre-existing draft output (so a leftover draft from a prior run is
    never re-consumed as if it were new source material). knowledge_base_
    path is None when KNOWLEDGE_BASE is a URL rather than a local path -
    there's nothing on disk to exclude in that case. Returns [] if the
    folder doesn't exist, which callers treat the same as "empty".
    """
    if not data_dir.is_dir():
        return []

    def _resolve(p: Path) -> Path:
        try:
            return p.resolve()
        except OSError:
            return p

    excluded = {_resolve(draft_path)}
    if knowledge_base_path is not None:
        excluded.add(_resolve(knowledge_base_path))
    files = []
    for p in sorted(data_dir.iterdir()):
        if not p.is_file() or p.name.startswith("."):
            continue
        if _resolve(p) in excluded:
            continue
        files.append(p)
    return files


def build_profile_prompt(existing_kb: dict, source_texts: dict) -> str:
    """
    Builds the system prompt for call_llm_update_profile(). Branches
    on whether an existing knowledge base was supplied: existing_kb is
    None for a from-scratch build (no KNOWLEDGE_BASE file yet), or a dict
    for a non-destructive update of one that already exists.
    """
    sources_block = "\n\n".join(
        f"=== SOURCE FILE: {name} ===\n{text}" for name, text in source_texts.items()
    )

    if existing_kb is not None:
        role_instructions = (
            "You are updating an EXISTING resume knowledge base JSON file "
            "with new information extracted from the source documents "
            "below. This update must be NON-DESTRUCTIVE: preserve every "
            "existing top-level section, field, and entry in the EXISTING "
            "KNOWLEDGE BASE exactly as it is, unless a source document "
            "gives new, more current, or corrected information for that "
            "exact item (for example a new job, a new skill, or an "
            "updated end date). Never remove, blank out, or shorten "
            "existing employers, skills, education entries, or any other "
            "existing content. Only ADD new entries where the source "
            "documents provide genuinely new information, and only "
            "MODIFY an existing entry when a source document clearly "
            "updates that specific fact."
        )
        base_block = "=== EXISTING KNOWLEDGE BASE (JSON) ===\n" + json.dumps(existing_kb, indent=2) + "\n\n"
    else:
        role_instructions = (
            "You are building a BRAND NEW resume knowledge base JSON file "
            "from scratch, using only the information present in the "
            "source documents below."
        )
        base_block = ""

    return (
        role_instructions + "\n\n"
        "Never fabricate information (employers, titles, dates, degrees, "
        "certifications, skills, or quantified metrics) that isn't "
        "present in the source documents (and, if given, the existing "
        "knowledge base). The output must be a single JSON object with "
        "AT LEAST these top-level keys: personal_info, education, "
        "skills, work_experience. Include any other section (e.g. "
        "summary_variants, cover_letter_building_blocks) that's clearly "
        "supported by the source material, matching the shape of the "
        "EXISTING KNOWLEDGE BASE where one is given.\n\n"
        + base_block
        + "=== SOURCE DOCUMENTS ===\n" + sources_block + "\n\n"
        "=== YOUR TASK ===\n"
        "Output ONLY the final JSON object - no commentary, no markdown "
        "code fences, no leading or trailing text. Start your response "
        "with '{' and end it with '}'."
    )


def validate_profile_draft(data: dict, existing_kb: dict) -> None:
    if not isinstance(data, dict):
        raise ValueError("Draft knowledge base is not a JSON object.")
    required = ("personal_info", "education", "skills", "work_experience")
    missing = [k for k in required if k not in data]
    if missing:
        raise ValueError(f"Draft knowledge base is missing required section(s): {missing}")
    if existing_kb is not None:
        dropped = [k for k in existing_kb if k not in data]
        if dropped:
            raise ValueError(
                f"Update dropped existing top-level section(s), which must stay non-destructive: {dropped}"
            )


def call_llm_update_profile(client: openai.OpenAI, existing_kb: dict, source_texts: dict) -> dict:
    system_prompt = build_profile_prompt(existing_kb, source_texts)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "Produce the JSON knowledge base now."},
    ]
    max_output_tokens = int(os.environ.get("LITELLM_MAX_TOKENS", "10000"))
    num_ctx = int(os.environ.get("OLLAMA_NUM_CTX", "16384"))

    est_input_tokens = len(system_prompt) // 4
    print(f"[profile] Prompt size estimate: ~{est_input_tokens} input tokens.", file=sys.stderr)

    last_error = None
    for attempt in range(2):
        try:
            response = client.chat.completions.create(
                model=MODEL, max_tokens=max_output_tokens, temperature=0.2, timeout=LLM_REQUEST_TIMEOUT,
                extra_body={"options": {"num_ctx": num_ctx}, "keep_alive": LITELLM_KEEP_ALIVE}, messages=messages,
            )
        except openai.APIConnectionError as e:
            print(f"[profile] Could not reach LiteLLM at {client.base_url}: {e}", file=sys.stderr)
            sys.exit(1)
        except openai.APIStatusError as e:
            print(f"[profile] LiteLLM returned an error (HTTP {e.status_code}): {e.message}", file=sys.stderr)
            sys.exit(1)

        raw = response.choices[0].message.content or ""
        try:
            data = json.loads(extract_json_object(raw))
            validate_profile_draft(data, existing_kb)
            return data
        except (ValueError, json.JSONDecodeError) as e:
            last_error = e
            print(f"[profile] Attempt {attempt + 1}: malformed output ({e}). Raw:\n{raw}\n", file=sys.stderr)
            messages.append({"role": "assistant", "content": raw})
            messages.append({
                "role": "user",
                "content": f"That response was invalid: {e}. Output ONLY the corrected, complete JSON object, following all the same rules.",
            })

    print(f"[profile] Model failed to produce a valid knowledge base after 2 attempts. Last error: {last_error}", file=sys.stderr)
    sys.exit(1)


def fetch_knowledge_base_json(url: str) -> dict:
    """
    Fetches and parses the knowledge base from a URL, for a KNOWLEDGE_BASE
    value that's http(s) instead of a local path - e.g. a raw file URL
    into a private repo, so the knowledge base can live outside this
    checkout entirely (see load_knowledge_base()). If KNOWLEDGE_BASE_
    URL_TOKEN is set, it's sent as an "Authorization: token <value>"
    header (GitHub's convention - works against both api.github.com and
    raw.githubusercontent.com with a personal access token, for a private
    repo). Raises ValueError, folding in the underlying error, on any
    connection, timeout, non-2xx-status, or invalid-JSON failure.
    """
    headers = {"User-Agent": "Mozilla/5.0 (compatible; resume-generator/1.0)"}
    token = os.environ.get("KNOWLEDGE_BASE_URL_TOKEN")
    if token:
        headers["Authorization"] = f"token {token}"

    try:
        response = httpx.get(url, timeout=30.0, follow_redirects=True, headers=headers)
        response.raise_for_status()
    except httpx.HTTPError as e:
        raise ValueError(f"Could not fetch KNOWLEDGE_BASE from {url!r}: {e}") from e

    try:
        return response.json()
    except ValueError as e:
        raise ValueError(f"KNOWLEDGE_BASE at {url!r} did not return valid JSON: {e}") from e


def load_knowledge_base(location: str) -> dict:
    """
    Loads the knowledge base from KNOWLEDGE_BASE, which may be either a
    local file path (the original behavior) or an http(s) URL. Raises
    ValueError (fetch/parse failure) for a URL, or the usual
    FileNotFoundError/json.JSONDecodeError for a local path.
    """
    if urllib.parse.urlparse(location).scheme in ("http", "https"):
        return fetch_knowledge_base_json(location)
    return json.loads(Path(location).read_text(encoding="utf-8"))


def generate_profile_draft(client: openai.OpenAI, s: dict) -> None:
    """
    Implements --generate profile: builds or non-destructively updates
    a profile.json (next to KNOWLEDGE_BASE) from whatever source
    files (pdf/txt/json/xml/docx) are sitting in the DATA folder, via
    LiteLLM, then removes the consumed source files - never the
    KNOWLEDGE_BASE file itself, and never the draft it just wrote.

    Per spec, this is a series of "nothing happens" short-circuits:
      * DATA env var not set at all -> nothing happens.
      * DATA folder has no source files (folder missing, or empty aside
        from the knowledge base / a stale draft) -> nothing happens,
        regardless of whether KNOWLEDGE_BASE exists.
      * Otherwise: KNOWLEDGE_BASE missing -> build a new draft from
        source files alone. KNOWLEDGE_BASE present -> non-destructively
        update it into the draft using the source files.

    If KNOWLEDGE_BASE is a URL (see load_knowledge_base()), there's
    nothing on disk to read as the "existing" knowledge base or to
    exclude from the DATA folder scan, so it's fetched instead - the
    update is still non-destructive to it, since the merged result is
    always written locally to KNOWLEDGE_BASE_DRAFT rather than back to
    the URL (this tool has no way to push to an arbitrary URL; promoting
    the draft back to wherever KNOWLEDGE_BASE lives is a manual step).
    """
    data_dir_setting = s.get("DATA")
    if not data_dir_setting:
        print("[profile] DATA is not set; skipping.", file=sys.stderr)
        return

    data_dir = Path(data_dir_setting)
    kb_location = s["KNOWLEDGE_BASE"]
    kb_is_url = urllib.parse.urlparse(kb_location).scheme in ("http", "https")
    kb_path = None if kb_is_url else Path(kb_location)

    def _read_existing_kb():
        if kb_is_url:
            return fetch_knowledge_base_json(kb_location)
        return json.loads(kb_path.read_text(encoding="utf-8")) if kb_path.is_file() else None

    # KNOWLEDGE_BASE_DRAFT is a naming template too (see render_filename),
    # not a raw path -- e.g. "data/{datetime.now}/profile.json" nests
    # each run's draft under its own timestamped subfolder, and a
    # template using {FirstName}/{LastName}/{Email} needs a name to pull
    # from. This speculative read is ONLY for that: any read/fetch/parse
    # failure here falls back to "" for those placeholders rather than
    # aborting outright, since an unreadable/malformed/unreachable
    # KNOWLEDGE_BASE shouldn't crash a run that -- once the source-file
    # scan below runs -- may well have turned out to be a no-op anyway.
    # If KNOWLEDGE_BASE really is broken and there ARE source files to
    # process, that surfaces properly below, where it's read again for
    # real.
    try:
        personal_info = (_read_existing_kb() or {}).get("personal_info", {})
    except (OSError, ValueError):
        personal_info = {}
    draft_path = Path(render_filename(
        s["KNOWLEDGE_BASE_DRAFT"], personal_info.get("full_name", ""), "", "json",
        email=personal_info.get("email", ""),
    ))

    source_files = build_source_file_list(data_dir, kb_path, draft_path)
    if not source_files:
        print(f"[profile] No source files found in {data_dir}/; skipping.", file=sys.stderr)
        return

    source_texts = {}
    for f in source_files:
        text = extract_text_from_source_file(f)
        if text.strip():
            source_texts[f.name] = text
    if not source_texts:
        print(f"[profile] Source files in {data_dir}/ had no extractable text; skipping.", file=sys.stderr)
        return

    try:
        existing_kb = _read_existing_kb()
    except (OSError, ValueError) as e:
        print(f"[profile] {e}", file=sys.stderr)
        sys.exit(1)

    draft = call_llm_update_profile(client, existing_kb, source_texts)
    ensure_parent_dir_exists(draft_path).write_text(json.dumps(draft, indent=2), encoding="utf-8")
    verb = "Updated" if existing_kb is not None else "Built"
    print(f"[profile] {verb} knowledge base at {draft_path}")
    if kb_is_url:
        print(
            f"[profile] KNOWLEDGE_BASE is a URL ({kb_location}); the draft above was NOT pushed there "
            "-- review it, then promote it back to that source yourself.",
            file=sys.stderr,
        )

    for f in source_files:
        try:
            f.unlink()
        except OSError as e:
            print(f"[profile] Could not remove consumed source file {f}: {e}", file=sys.stderr)
    print(f"[profile] Removed {len(source_files)} consumed source file(s) from {data_dir}/")


# --- --analyze: job-fit analysis against a job description ----------------

# Every key the model's top-level JSON response must contain for
# validate_job_fit_analysis() to accept it.
REQUIRED_ANALYSIS_KEYS = {"fit_assessments", "overall_summary", "upskill_resources"}

# Every key each entry of fit_assessments must contain. "matched_
# qualifications" is intentionally NOT required - it's useful context
# for the reader but, unlike the others, isn't something the feature spec
# asked for, so a model that omits it shouldn't burn a retry.
REQUIRED_ASSESSMENT_KEYS = {"list_label", "fit_percentage", "assessment_summary", "missing_qualifications"}

# Keys each entry of upskill_resources must have. resource_url and is_free
# are deliberately NOT required: a model that (correctly) doesn't know a
# real URL for a given resource should say so via resource_name alone
# rather than being forced to invent one to pass validation (see
# never_fabricate in output_rules, and build_job_fit_prompt below).
REQUIRED_RESOURCE_KEYS = {"missing_item", "resource_name"}


def _strip_html(markup: str) -> str:
    """
    Reduces an HTML document to plain text for _fetch_job_description_
    from_url: drops <script>/<style> blocks entirely, strips remaining
    tags, unescapes entities, and collapses excess whitespace. Not a
    real HTML parser - just enough to pull readable text out of a job
    posting page without adding a new dependency (the project's
    requirements.txt has no HTML parser; beautifulsoup4 is test-only,
    see requirements-test.txt).
    """
    markup = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", markup)
    text = re.sub(r"(?s)<[^>]+>", " ", markup)
    text = html.unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()


def _fetch_job_description_from_url(url: str) -> str:
    """
    Best-effort fetch of job description text from a URL, for --analyze.
    If the response looks like HTML, it's reduced to plain text via
    _strip_html; otherwise the response body is used as-is (e.g. an API
    endpoint returning the posting as plain text or JSON). Raises
    ValueError, folding in the underlying error, on any connection,
    timeout, or non-2xx-status failure.
    """
    try:
        response = httpx.get(
            url, timeout=30.0, follow_redirects=True,
            headers={"User-Agent": "Mozilla/5.0 (compatible; resume-generator/1.0)"},
        )
        response.raise_for_status()
    except httpx.HTTPError as e:
        raise ValueError(f"Could not fetch job description from {url!r}: {e}") from e

    body = response.text
    content_type = response.headers.get("content-type", "")
    if "html" in content_type.lower() or "<html" in body[:1000].lower():
        return _strip_html(body)
    return body.strip()


def resolve_job_description(raw: str) -> str:
    """
    Turns --analyze's raw CLI value into job description text. raw is
    interpreted, in order:
      1. A URL (http:// or https:// scheme) - fetched, and reduced to
         plain text if it looks like HTML (see
         _fetch_job_description_from_url).
      2. A path to an existing local file - its text is extracted
         (reusing extract_text_from_source_file, so pdf/docx/txt/md/json/
         xml all work, same as a DATA-folder source file).
      3. Otherwise, raw itself: the job description text pasted directly
         on the command line.
    Raises ValueError if the result is empty (an unreachable/erroring
    URL, an existing-but-empty-or-unreadable file, or an all-whitespace
    literal value).
    """
    if urllib.parse.urlparse(raw).scheme in ("http", "https"):
        text = _fetch_job_description_from_url(raw)
        if not text:
            raise ValueError(f"--analyze URL {raw!r} returned no extractable text.")
        return text

    path = Path(raw)
    if path.is_file():
        text = extract_text_from_source_file(path)
        if not text.strip():
            raise ValueError(f"--analyze file {raw!r} had no extractable text.")
        return text.strip()

    text = raw.strip()
    if not text:
        raise ValueError("--analyze's job description value was empty.")
    return text


def build_job_fit_context(kb: dict) -> dict:
    """
    Trims the full knowledge base down to what the job-fit analysis needs:
    unlike build_baseline_context (which is scoped to ONE resume variant),
    this pulls skills and work-experience bullets across BOTH variants,
    since the model needs the candidate's full skill set to judge fit
    against an arbitrary job description, not just what one resume
    variant would show.
    """
    skills = [
        {"category": CATEGORY_LABELS.get(k, k.replace("_", " ").title()), "items": v}
        for k, v in kb["skills"].items() if isinstance(v, list)
    ]
    work_experience = [
        {
            "titles": job["title_by_variant"],
            "company": job["company"],
            "date_range": f"{job['start_date']} - {job['end_date']}",
            "technologies": job.get("technologies", []),
            "bullets": [b["text"] for b in job["bullets"] if not b["id"].endswith(("_alt", "_variant"))],
        }
        for job in kb["work_experience"]
    ]
    education = [
        {"degree": ed["degree"], "institution": ed["institution"], "date": ed["graduation date"]}
        for ed in kb["education"]
    ]
    return {
        "summaries": kb["summary_variants"],
        "education": education,
        "skills": skills,
        "work_experience": work_experience,
    }


def build_job_fit_prompt(kb: dict, job_description: str, prompt_template_text: str) -> str:
    """
    Fills prompt_template_text (the contents of ANALYSIS_PROMPT_TEMPLATE,
    e.g. ANALYSIS_PROMPT.template.txt - a string.Template using
    $output_rules/$candidate_data/$job_description placeholders) in with
    this run's actual data. string.Template (not str.format) is used
    deliberately: the template's JSON schema example is full of literal
    { } characters that would otherwise all need doubling up to escape
    them from str.format.
    """
    context = build_job_fit_context(kb)
    rules = kb["meta"]["output_rules"]
    template = string.Template(prompt_template_text)
    return template.substitute(
        output_rules=json.dumps({"never_fabricate": rules["never_fabricate"]}, indent=2),
        candidate_data=json.dumps(context, indent=2),
        job_description=job_description,
    )


def validate_job_fit_analysis(data: dict) -> None:
    """Raises ValueError on any structural problem in the model's job-fit
    JSON, which call_llm_analyze_fit uses to trigger a corrective retry."""
    if not isinstance(data, dict):
        raise ValueError(f"Expected a JSON object, got {type(data).__name__}.")
    missing_keys = REQUIRED_ANALYSIS_KEYS - set(data.keys())
    if missing_keys:
        raise ValueError(f"Missing required key(s): {sorted(missing_keys)}")

    if not isinstance(data["overall_summary"], str) or not data["overall_summary"].strip():
        raise ValueError("overall_summary must be a non-empty string.")

    assessments = data["fit_assessments"]
    if not isinstance(assessments, list) or not assessments:
        raise ValueError("fit_assessments must be a non-empty list.")
    for i, assessment in enumerate(assessments):
        if not isinstance(assessment, dict):
            raise ValueError(f"fit_assessments[{i}] must be an object, got {type(assessment).__name__}.")
        missing = REQUIRED_ASSESSMENT_KEYS - set(assessment.keys())
        if missing:
            raise ValueError(f"fit_assessments[{i}] is missing required key(s): {sorted(missing)}")

        if not isinstance(assessment["list_label"], str) or not assessment["list_label"].strip():
            raise ValueError(f"fit_assessments[{i}].list_label must be a non-empty string.")

        pct = assessment["fit_percentage"]
        if isinstance(pct, bool) or not isinstance(pct, (int, float)):
            raise ValueError(f"fit_assessments[{i}].fit_percentage must be a number, got {pct!r}.")
        if not (0 <= pct <= 100):
            raise ValueError(f"fit_assessments[{i}].fit_percentage must be between 0 and 100, got {pct!r}.")

        if not isinstance(assessment["assessment_summary"], str) or not assessment["assessment_summary"].strip():
            raise ValueError(f"fit_assessments[{i}].assessment_summary must be a non-empty string.")

        if not isinstance(assessment["missing_qualifications"], list) or not all(
            isinstance(x, str) for x in assessment["missing_qualifications"]
        ):
            raise ValueError(f"fit_assessments[{i}].missing_qualifications must be a list of strings.")

        if "matched_qualifications" in assessment and (
            not isinstance(assessment["matched_qualifications"], list)
            or not all(isinstance(x, str) for x in assessment["matched_qualifications"])
        ):
            raise ValueError(f"fit_assessments[{i}].matched_qualifications must be a list of strings.")

    resources = data["upskill_resources"]
    if not isinstance(resources, list):
        raise ValueError("upskill_resources must be a list.")
    for i, res in enumerate(resources):
        if not isinstance(res, dict):
            raise ValueError(f"upskill_resources[{i}] must be an object, got {type(res).__name__}.")
        missing = REQUIRED_RESOURCE_KEYS - set(res.keys())
        if missing:
            raise ValueError(f"upskill_resources[{i}] is missing required key(s): {sorted(missing)}")


def call_llm_analyze_fit(client: openai.OpenAI, kb: dict, job_description: str, prompt_template_text: str) -> dict:
    system_prompt = build_job_fit_prompt(kb, job_description, prompt_template_text)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "Analyze the fit now."},
    ]
    max_output_tokens = int(os.environ.get("LITELLM_MAX_TOKENS", "10000"))
    num_ctx = int(os.environ.get("OLLAMA_NUM_CTX", "16384"))

    est_input_tokens = len(system_prompt) // 4
    print(f"[analyze] Prompt size estimate: ~{est_input_tokens} input tokens.", file=sys.stderr)

    last_error = None
    for attempt in range(2):
        try:
            try:
                response = client.chat.completions.create(
                    model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=LLM_REQUEST_TIMEOUT,
                    response_format={"type": "json_object"},
                    extra_body={"options": {"num_ctx": num_ctx}, "keep_alive": LITELLM_KEEP_ALIVE}, messages=messages,
                )
            except openai.BadRequestError:
                response = client.chat.completions.create(
                    model=MODEL, max_tokens=max_output_tokens, temperature=0.3, timeout=LLM_REQUEST_TIMEOUT,
                    extra_body={"options": {"num_ctx": num_ctx}, "keep_alive": LITELLM_KEEP_ALIVE}, messages=messages,
                )
        except openai.APIConnectionError as e:
            print(f"[analyze] Could not reach LiteLLM at {client.base_url}: {e}", file=sys.stderr)
            sys.exit(1)
        except openai.APIStatusError as e:
            print(f"[analyze] LiteLLM returned an error (HTTP {e.status_code}): {e.message}", file=sys.stderr)
            sys.exit(1)

        raw = response.choices[0].message.content or ""
        try:
            analysis = json.loads(extract_json_object(raw))
            validate_job_fit_analysis(analysis)
            return analysis
        except (ValueError, json.JSONDecodeError) as e:
            last_error = e
            print(f"[analyze] Attempt {attempt + 1}: malformed analysis output ({e}). Raw:\n{raw}\n", file=sys.stderr)
            messages.append({"role": "assistant", "content": raw})
            messages.append({
                "role": "user",
                "content": f"That response was invalid: {e}. Respond again with ONLY the JSON object described earlier.",
            })

    print(f"[analyze] Model failed to produce a valid analysis after 2 attempts. Last error: {last_error}", file=sys.stderr)
    sys.exit(1)


def render_job_fit_analysis_md(analysis: dict) -> str:
    """Renders call_llm_analyze_fit()'s parsed dict as a human-readable
    markdown report. analysis["fit_assessments"] always has at least one
    entry: exactly one when the job description didn't separate its
    qualifications into distinct lists (e.g. "Required" vs "Preferred"),
    or one per list when it did - see ANALYSIS_PROMPT_TEMPLATE. The
    single-list case is rendered flat (one fit score, one summary, one
    matched/missing pair) rather than nesting it under a redundant
    per-list breakdown."""
    assessments = analysis["fit_assessments"]
    lines = ["# Job Fit Analysis", ""]

    if len(assessments) == 1:
        a = assessments[0]
        lines += [f"**Fit score: {a['fit_percentage']}%**", "", "## Summary", "", a["assessment_summary"], ""]
        lines += _render_matched_and_missing_md(a, heading_level="##")
    else:
        lines += ["## Fit Scores", ""]
        lines += [f"- **{a['list_label']}**: {a['fit_percentage']}%" for a in assessments]
        lines += ["", "## Summary", "", analysis["overall_summary"], ""]
        for a in assessments:
            lines += [f"## {a['list_label']}", "", f"**Fit score: {a['fit_percentage']}%**", "", a["assessment_summary"], ""]
            lines += _render_matched_and_missing_md(a, heading_level="###")

    resources = analysis["upskill_resources"]
    lines += ["## Suggested Resources to Close the Gaps", ""]
    if resources:
        for r in resources:
            free_tag = " (Free)" if r.get("is_free") else ""
            type_tag = f" [{r['resource_type']}]" if r.get("resource_type") else ""
            name = r["resource_name"]
            if r.get("resource_url"):
                name = f"[{name}]({r['resource_url']})"
            lines.append(f"- **{r['missing_item']}**: {name}{type_tag}{free_tag}")
    else:
        lines.append("No resources suggested.")

    return "\n".join(lines) + "\n"


def _render_matched_and_missing_md(assessment: dict, heading_level: str) -> list:
    """Shared by render_job_fit_analysis_md's single- and multi-assessment
    branches: renders one assessment's matched/missing qualifications as
    markdown lines, at the given heading level ("##" or "###")."""
    lines = []
    matched = assessment.get("matched_qualifications") or []
    if matched:
        lines += [f"{heading_level} Matched Qualifications", ""]
        lines += [f"- {m}" for m in matched]
        lines.append("")

    missing = assessment["missing_qualifications"]
    lines += [f"{heading_level} Missing Qualifications", ""]
    if missing:
        lines += [f"- {m}" for m in missing]
    else:
        lines.append("None found - the candidate data covers every requirement in this list.")
    lines.append("")
    return lines


def load_file_location_settings() -> dict:
    return {
        # Not strictly a file location like the rest of this dict, but
        # lives here for the same reason the OUTPUT_REPO* settings do:
        # one place main() reads every env-configurable setting from, so
        # it's re-read (and testable) fresh per run/test rather than
        # frozen at import time.
        #
        # Comma-separated resume/cover-letter variants to generate, in
        # order - e.g. "SDE,SDET" (the default) or "SDE,SDET,SRE". Each
        # entry becomes {JobAcronym} in a naming template, a key looked
        # up in profile.json's per-variant fields (summary_variants,
        # title_by_variant, etc. - see build_baseline_context), and,
        # unless it's "SDE" or "SDET", falls back to a generic
        # "CORE <VARIANT> SKILLS" heading (see SKILLS_HEADING_BY_VARIANT)
        # rather than failing outright.
        "VARIANTS": [v.strip() for v in os.environ.get("VARIANTS", "SDE,SDET").split(",") if v.strip()],
        "OUTPUT_FOLDER": os.environ.get("OUTPUT_FOLDER", "generated"),
        "KNOWLEDGE_BASE": os.environ.get("KNOWLEDGE_BASE", "data/profile.json"),
        "KNOWLEDGE_BASE_DRAFT": os.environ.get("KNOWLEDGE_BASE_DRAFT", "data/profile.json"),
        "DATA": os.environ.get("DATA"),
        "README_TEMPLATE": os.environ.get("README_TEMPLATE", "README.template.md"),
        "README_OUTPUT": os.environ.get("README_OUTPUT", "README.md"),
        "RESUME_TEMPLATE": os.environ.get("RESUME_TEMPLATE", "RESUME.template.md"),
        "RESUME_NAMING_TEMPLATE": os.environ.get(
            "RESUME_NAMING_TEMPLATE", "{FirstName} {LastName} Resume ({JobAcronym}).{Extension}"
        ),
        "COVERLETTER_NAMING_TEMPLATE": os.environ.get(
            "COVERLETTER_NAMING_TEMPLATE", "{FirstName} {LastName} Cover Letter ({JobAcronym}).{Extension}"
        ),
        "ANALYSIS_PROMPT_TEMPLATE": os.environ.get("ANALYSIS_PROMPT_TEMPLATE", "ANALYSIS_PROMPT.template.txt"),
        # OUTPUT_REPO and friends: see sync_output_repo() and
        # commit_and_push_output_repo(). Unset OUTPUT_REPO (the default)
        # means "no change" - OUTPUT_FOLDER is written into this checkout,
        # exactly like before this feature existed.
        "OUTPUT_REPO": os.environ.get("OUTPUT_REPO", ""),
        "OUTPUT_REPO_BRANCH": os.environ.get("OUTPUT_REPO_BRANCH", ""),
        "OUTPUT_REPO_USER": os.environ.get("OUTPUT_REPO_USER", ""),
        "OUTPUT_REPO_TOKEN": os.environ.get("OUTPUT_REPO_TOKEN", ""),
        "OUTPUT_REPO_CLONE_DIR": os.environ.get("OUTPUT_REPO_CLONE_DIR", ".output-repo"),
        "OUTPUT_REPO_AUTHOR_NAME": os.environ.get("OUTPUT_REPO_AUTHOR_NAME", "Dennis Jay Dole"),
        "OUTPUT_REPO_AUTHOR_EMAIL": os.environ.get(
            "OUTPUT_REPO_AUTHOR_EMAIL", "Dennis.Dole+resume-generator@djdole.net"
        ),
        "OUTPUT_REPO_COMMIT_MESSAGE": os.environ.get(
            "OUTPUT_REPO_COMMIT_MESSAGE", "Regenerate resumes/cover letters ({datetime.now})"
        ),
        "OUTPUT_REPO_PUSH": os.environ.get("OUTPUT_REPO_PUSH", "true").strip().lower() not in ("false", "0", "no"),
    }


class _NowPlaceholder(str):
    """
    A str subclass wrapping a single "now" timestamp, so a naming-
    template placeholder can be used bare (renders as a filesystem-safe
    default format) OR with dotted attribute access for an individual
    component -- both from the exact same `datetime.now()` call, so a
    template using several of these in one path (e.g. a year folder and
    a day file) can't straddle a rollover between them.

    Bare {datetime.now} -> "YYYY-MM-DD_HHMMSS" (sortable, no ':' or ' ',
    safe as a path segment on every OS this runs on).
    {datetime.now.year}/{.month}/{.day}/{.hour}/{.minute}/{.second} ->
    the real underlying int (so e.g. {datetime.now.month:02d} zero-pads
    via a normal str.format spec) -- and any other real datetime.datetime
    attribute/method (.strftime, .date, .weekday, ...) is reachable the
    same way, since unresolved attribute lookups fall through to it.
    """

    def __new__(cls, dt: datetime.datetime):
        obj = super().__new__(cls, dt.strftime("%Y-%m-%d_%H%M%S"))
        obj._dt = dt
        return obj

    def __getattr__(self, name):
        return getattr(self._dt, name)


def render_filename(
    naming_template: str, full_name: str, job_acronym: str, extension: str, email: str = "",
) -> str:
    """
    Fills a naming/output-path template -- RESUME_NAMING_TEMPLATE,
    COVERLETTER_NAMING_TEMPLATE, or KNOWLEDGE_BASE_DRAFT -- with this
    run's actual values. See USAGE.md's "Naming template placeholders"
    section for the full list; in short:
      {FirstName}/{LastName} - full_name's first and last whitespace-
        separated token (a middle name/initial is dropped). Both render
        as "" if full_name is falsy (e.g. KNOWLEDGE_BASE_DRAFT during a
        from-scratch profile build, before any name is known yet).
      {JobAcronym} - job_acronym verbatim (e.g. "SDE"/"SDET"); "" where
        not applicable (KNOWLEDGE_BASE_DRAFT isn't per-variant).
      {Extension} - extension verbatim (e.g. "pdf").
      {Email} - email verbatim; "" if not given.
      {datetime.now...} - see _NowPlaceholder.

    A template MAY contain "/" to nest the result under a subfolder
    (e.g. "{JobAcronym}/{FirstName} {LastName} Resume.{Extension}") --
    render_filename only fills in the placeholders; creating that
    subfolder before writing anything into it is the caller's job, via
    ensure_parent_dir_exists().
    """
    parts = (full_name or "").split()
    first_name = parts[0] if parts else ""
    last_name = parts[-1] if parts else ""
    now_ns = types.SimpleNamespace(now=_NowPlaceholder(datetime.datetime.now()))
    return naming_template.format(
        FirstName=first_name, LastName=last_name, JobAcronym=job_acronym, Extension=extension,
        Email=email or "", datetime=now_ns,
    )


def ensure_parent_dir_exists(path: Path) -> Path:
    """
    Creates path's parent directory (and any missing intermediate
    directories) if it doesn't already exist yet, then returns path
    unchanged -- so a write call site can just be
    `ensure_parent_dir_exists(some_path).write_text(...)`, or pass the
    wrapped path straight to a renderer that writes to it internally
    (render_resume_pdf, render_resume_docx, render_cover_letter_pdf,
    render_cover_letter_docx). This is what makes a naming template like
    "{JobAcronym}/{FirstName} {LastName} Resume.{Extension}" work at all:
    reportlab/python-docx/plain file writes all fail outright if the
    directory they're writing into doesn't exist yet, and naming
    templates are free to put a placeholder before the LAST '/' to
    request a subfolder that may not exist on disk until this runs.
    """
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _run_git(args: list, cwd, env: dict = None) -> str:
    """
    Runs `git <args>` in cwd and returns stdout. Raises RuntimeError
    (folding in git's own stderr) on a non-zero exit, so an OUTPUT_REPO
    problem (bad URL, auth failure, network error, non-fast-forward
    push, ...) surfaces as a clear error instead of silently no-op'ing.
    """
    result = subprocess.run(["git", *args], cwd=cwd, env=env, capture_output=True, text=True)
    if result.returncode != 0:
        raise RuntimeError(f"git {' '.join(args)} (in {cwd}) failed: {result.stderr.strip()}")
    return result.stdout


def _inject_repo_token(repo_url: str, user: str,token: str) -> str:
    """
    Embeds OUTPUT_REPO_TOKEN into an https:// OUTPUT_REPO URL as an HTTP
    Basic auth credential, so clone/fetch/push work non-interactively
    (e.g. in CI) without SSH key setup -- GitHub/GitLab/etc. all accept a
    personal access token as the password with any non-empty username;
    "x-access-token" is used here, GitHub's own convention for this.
    Leaves ssh:// URLs and local paths untouched (there's no equivalent
    embedding for those -- use your normal SSH key/ssh-agent setup
    instead), and leaves the URL untouched entirely when no token is set.
    """
    if not token or not repo_url.startswith(("http://", "https://")):
        return repo_url
    parsed = urllib.parse.urlsplit(repo_url)
    usr = user or "x-access-token"
    netloc = f"{usr}:{token}@{parsed.netloc}"
    return urllib.parse.urlunsplit((parsed.scheme, netloc, parsed.path, parsed.query, parsed.fragment))


def _resolve_output_repo_target_ref(clone_dir: Path, branch: str):
    """
    Returns the origin ref sync_output_repo() should hard-reset the clone
    to, or None if that ref doesn't exist yet -- which just means the
    remote has no commits reachable from it yet (e.g. this is the very
    first run ever against a freshly created, still-empty OUTPUT_REPO,
    or a retry after a previous run crashed before
    commit_and_push_output_repo() got to push anything). In that case
    there's nothing to reset TO, and the caller falls back to only
    cleaning up untracked leftovers instead.
    """
    if branch:
        candidate = f"origin/{branch}"
    else:
        head = subprocess.run(
            ["git", "symbolic-ref", "-q", "refs/remotes/origin/HEAD"], cwd=clone_dir, capture_output=True, text=True,
        )
        if head.returncode != 0:
            return None
        candidate = head.stdout.strip().removeprefix("refs/remotes/")
    exists = subprocess.run(
        ["git", "rev-parse", "--verify", "-q", candidate], cwd=clone_dir, capture_output=True,
    ).returncode == 0
    return candidate if exists else None


def sync_output_repo(s: dict) -> Path:
    """
    Ensures OUTPUT_REPO is available locally at OUTPUT_REPO_CLONE_DIR --
    cloning it there if this is the first run, or bringing an existing
    clone from a previous run up to date with origin otherwise -- and
    returns that clone's path. Called once per run, before any
    generation happens, whenever OUTPUT_REPO is set; see
    commit_and_push_output_repo() for the other half, after generation.

    Reusing OUTPUT_REPO_CLONE_DIR across runs (rather than a fresh clone
    every time) means repeated runs only fetch what changed. The clone is
    always reset to exactly match origin/<branch> before returning --
    any uncommitted changes or untracked files left over from an
    interrupted previous run (e.g. this script crashing after writing
    output but before commit_and_push_output_repo() ran) are discarded,
    so every run starts from a known-clean state instead of silently
    layering new output on top of leftover partial output.
    """
    clone_dir = Path(s["OUTPUT_REPO_CLONE_DIR"])
    repo_url = _inject_repo_token(s["OUTPUT_REPO"], s["OUTPUT_REPO_USER"], s["OUTPUT_REPO_TOKEN"])
    branch = s["OUTPUT_REPO_BRANCH"]

    if (clone_dir / ".git").is_dir():
        _run_git(["remote", "set-url", "origin", repo_url], clone_dir)
        _run_git(["fetch", "origin"], clone_dir)
        if not branch:
            # (Re-)detect origin's default branch. Needed whenever the
            # clone was originally made from an empty remote (no default
            # branch existed yet to record) and that remote has since
            # gained one -- e.g. this generator's own first successful
            # commit_and_push_output_repo() call. Failure (remote is
            # still empty) is fine and expected; _resolve_output_repo_
            # target_ref then falls through to the clean-only path below.
            subprocess.run(["git", "remote", "set-head", "origin", "-a"], cwd=clone_dir, capture_output=True)
        target = _resolve_output_repo_target_ref(clone_dir, branch)
        if target:
            if branch:
                _run_git(["checkout", "-B", branch, target], clone_dir)
            _run_git(["reset", "--hard", target], clone_dir)
        _run_git(["clean", "-fd"], clone_dir)
    else:
        if clone_dir.exists() and any(clone_dir.iterdir()):
            raise RuntimeError(
                f"OUTPUT_REPO_CLONE_DIR '{clone_dir}' already exists and isn't a git clone "
                "-- remove it, or point OUTPUT_REPO_CLONE_DIR somewhere else, and re-run."
            )
        clone_dir.parent.mkdir(parents=True, exist_ok=True)
        args = ["clone", repo_url, str(clone_dir)]
        if branch:
            args[1:1] = ["--branch", branch]
        _run_git(args, clone_dir.parent)

    return clone_dir


def commit_and_push_output_repo(s: dict, clone_dir: Path) -> bool:
    """
    Stages every change under clone_dir (the OUTPUT_REPO clone) and, if
    there's anything new to check in, commits it with
    OUTPUT_REPO_COMMIT_MESSAGE (a naming template -- see render_filename
    -- though only its {datetime.now...} placeholders make sense here,
    since a commit isn't per-resume-variant) and, unless OUTPUT_REPO_PUSH
    is false, pushes it to OUTPUT_REPO_BRANCH (or whatever branch is
    currently checked out, if that's unset). Returns True if a commit
    was made, False if the working tree already matched what's
    committed (nothing new to check in this run).
    """
    _run_git(["add", "-A"], clone_dir)
    if not _run_git(["status", "--porcelain"], clone_dir).strip():
        return False

    commit_env = {
        **os.environ,
        "GIT_AUTHOR_NAME": s["OUTPUT_REPO_AUTHOR_NAME"],
        "GIT_AUTHOR_EMAIL": s["OUTPUT_REPO_AUTHOR_EMAIL"],
        "GIT_COMMITTER_NAME": s["OUTPUT_REPO_AUTHOR_NAME"],
        "GIT_COMMITTER_EMAIL": s["OUTPUT_REPO_AUTHOR_EMAIL"],
    }
    message = render_filename(s["OUTPUT_REPO_COMMIT_MESSAGE"], full_name="", job_acronym="", extension="")
    _run_git(["commit", "-m", message], clone_dir, env=commit_env)

    if s["OUTPUT_REPO_PUSH"]:
        _run_git(["push", "origin", f"HEAD:{s['OUTPUT_REPO_BRANCH']}" if s["OUTPUT_REPO_BRANCH"] else "HEAD"], clone_dir)

    return True


def main(argv=None):
    parser = build_arg_parser()
    # argv is None (the default) whenever main() is called directly rather
    # than via the __main__ block below - e.g. from tests - in which
    # case there are no CLI args to parse (as opposed to argparse's own
    # default of falling back to sys.argv, which would pick up whatever
    # unrelated args the calling process - e.g. pytest - was invoked
    # with).
    args = parser.parse_args(argv if argv is not None else [])

    # --generate's "omitted entirely" default (generate everything) is
    # meant for the tool's primary, no-flags-at-all workflow. --analyze is
    # a separate, deliberately-opted-into action - someone running
    # `--analyze "..."` on its own wants ONLY the analysis, not to also
    # silently kick off a full resume/cover_letter/readme run. So that
    # default only applies when --analyze wasn't requested; an explicit
    # --generate (even alongside --analyze) always takes precedence over
    # both of those.
    if args.generate is None:
        targets = set() if args.analyze else set(ALL_TARGETS)
    else:
        targets = parse_generate_targets(args.generate)
        if not targets and not args.analyze:
            print("--generate was supplied with no value(s); nothing to generate.")
            return

    s = load_file_location_settings()

    # One client (and its underlying connection pool) for every LLM call
    # in this run, rather than a fresh one per call site.
    client = build_llm_client()

    # Only read KNOWLEDGE_BASE up front if a target actually needs it as
    # input. --generate profile has its own, separate rules about
    # whether KNOWLEDGE_BASE needs to exist yet (it may legitimately not),
    # so it reads it lazily, itself, inside generate_profile_draft().
    kb = full_name = None
    if {"resume", "cover_letter", "readme"} & targets or args.analyze:
        try:
            kb = load_knowledge_base(s["KNOWLEDGE_BASE"])
        except (OSError, ValueError) as e:
            print(f"[KNOWLEDGE_BASE] {e}", file=sys.stderr)
            sys.exit(1)
        full_name = kb["personal_info"]["full_name"]

    # Non-None only when OUTPUT_REPO is set and this run is actually
    # generating something OUTPUT_REPO covers (resume/cover_letter/
    # readme) - signals, after those blocks, whether
    # commit_and_push_output_repo() needs to run. Synced once up front,
    # rather than separately for the resume/cover_letter block and the
    # readme block below, so both write into the exact same clone/commit.
    output_repo_clone_dir = None
    if s["OUTPUT_REPO"] and ({"resume", "cover_letter", "readme"} & targets):
        try:
            output_repo_clone_dir = sync_output_repo(s)
        except RuntimeError as e:
            print(f"[OUTPUT_REPO] {e}", file=sys.stderr)
            sys.exit(1)

    if "resume" in targets or "cover_letter" in targets:
        out_dir = (
            output_repo_clone_dir / s["OUTPUT_FOLDER"] if output_repo_clone_dir is not None
            else Path(s["OUTPUT_FOLDER"])
        )
        out_dir.mkdir(parents=True, exist_ok=True)
        resume_template_text = (
            Path(s["RESUME_TEMPLATE"]).read_text(encoding="utf-8") if "resume" in targets else None
        )

        for variant in s["VARIANTS"]:
            def resume_path(ext: str) -> Path:
                return ensure_parent_dir_exists(out_dir / render_filename(s["RESUME_NAMING_TEMPLATE"], full_name, variant, ext))

            def cl_path(ext: str) -> Path:
                return ensure_parent_dir_exists(out_dir / render_filename(s["COVERLETTER_NAMING_TEMPLATE"], full_name, variant, ext))

            if "resume" in targets:
                r = call_llm_fill_resume(client, kb, variant, resume_template_text)

                resume_path("json").write_text(json.dumps(r, indent=2), encoding="utf-8")
                resume_path("txt").write_text(render_resume_txt(r), encoding="utf-8")
                resume_path("md").write_text(render_resume_md(r), encoding="utf-8")

                pages, body_pt = render_resume_pdf(r, resume_path("pdf"))
                if pages > 2:
                    print(f"WARNING: {variant} resume rendered at {pages} pages even at the smallest tier ({body_pt}pt).")
                else:
                    print(f"{variant} resume: {pages} page(s) at {body_pt}pt.")
                render_resume_docx(r, resume_path("docx"), body_pt=body_pt)

            if "cover_letter" in targets:
                cl = call_llm_cover_letter(client, kb, variant)

                cl_path("txt").write_text(render_cover_letter_txt(cl), encoding="utf-8")
                render_cover_letter_docx(cl, cl_path("docx"))
                render_cover_letter_pdf(cl, cl_path("pdf"))

    if "readme" in targets:
        # Separate LLM call, template-driven: fills README_TEMPLATE using
        # profile.json, rather than reusing the resume call above. When
        # OUTPUT_REPO is set, this is written into that same clone (and
        # so committed/pushed alongside the resume/cover letter output
        # below) rather than into this checkout - README_OUTPUT is a
        # generated file like any other target here.
        template_text = Path(s["README_TEMPLATE"]).read_text(encoding="utf-8")
        readme_markdown = call_llm_readme(client, kb, template_text)
        readme_path = (
            output_repo_clone_dir / s["README_OUTPUT"] if output_repo_clone_dir is not None
            else Path(s["README_OUTPUT"])
        )
        ensure_parent_dir_exists(readme_path).write_text(readme_markdown, encoding="utf-8")
        print(f"Wrote GitHub profile README to {readme_path}")

    if "profile" in targets:
        generate_profile_draft(client, s)

    if args.analyze:
        try:
            job_description = resolve_job_description(args.analyze)
        except ValueError as e:
            print(f"[analyze] {e}", file=sys.stderr)
            sys.exit(1)

        prompt_template_text = Path(s["ANALYSIS_PROMPT_TEMPLATE"]).read_text(encoding="utf-8")
        analysis = call_llm_analyze_fit(client, kb, job_description, prompt_template_text)

        rendered_job_fit_analysis_report = render_job_fit_analysis_md(analysis)
        print(rendered_job_fit_analysis_report)

    summary_parts = []
    if "resume" in targets:
        summary_parts.append(f"{len(s['VARIANTS'])} resume variant(s) (5 formats)")
    if "cover_letter" in targets:
        summary_parts.append("cover letters (3 formats)")
    if summary_parts:
        dest = f"{output_repo_clone_dir}/{s['OUTPUT_FOLDER']}" if output_repo_clone_dir is not None else s["OUTPUT_FOLDER"]
        print(f"Wrote {' + '.join(summary_parts)} to {dest}/")

    if output_repo_clone_dir is not None:
        try:
            committed = commit_and_push_output_repo(s, output_repo_clone_dir)
        except RuntimeError as e:
            print(f"[OUTPUT_REPO] {e}", file=sys.stderr)
            sys.exit(1)
        if committed:
            verb = "Committed and pushed" if s["OUTPUT_REPO_PUSH"] else "Committed (not pushed - OUTPUT_REPO_PUSH=false)"
            print(f"{verb} generated files to {s['OUTPUT_REPO']}.")
        else:
            print(f"No changes to check into {s['OUTPUT_REPO']} (output already up to date).")


if __name__ == "__main__":
    main(sys.argv[1:])